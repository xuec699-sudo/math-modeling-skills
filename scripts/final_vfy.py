# -*- coding: utf-8 -*-
import sys, io, zipfile, re, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

path = r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx"
doc = Document(path)

# Check key changes are present
checks = {
    "CV gap discussion (6.3x)": False,
    "systematic drift": False,
    "conservative estimate R²≈0.91": False,
    "hyperparameter config": False,
    "grid vs triangular": False,
    "三角排列": False,
    "[10] Wolanin": False,
    "[11] 张宝忠": False,
    "[12] Chen T": False,
}

for para in doc.paragraphs:
    t = para.text
    if "6.3倍" in t or "6.3" in t: checks["CV gap discussion (6.3x)"] = True
    if "分布漂移" in t or "systematic drift" in t: checks["systematic drift"] = True
    if "R²≈0.91" in t or "R²≈0.91" in t: checks["conservative estimate R²≈0.91"] = True
    if "超参数配置" in t or "200棵树" in t: checks["hyperparameter config"] = True
    if "三角" in t or "triangular" in t.lower(): checks["三角排列"] = True

for para in doc.paragraphs:
    t = para.text.strip()
    if t.startswith('[10]'): checks["[10] Wolanin"] = True
    if t.startswith('[11]'): checks["[11] 张宝忠"] = True
    if t.startswith('[12]'): checks["[12] Chen T"] = True

# ZIP check
with zipfile.ZipFile(path, "r") as z:
    xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    images = [n for n in z.namelist() if "media" in n and n.endswith('.png')]
    omml = len(re.findall(r"<m:oMath", xml))
    raw_dollar = len(re.findall(r"<w:t[^>]*>\$", xml))

print("=" * 50)
print("  FINAL VERIFICATION")
print("=" * 50)
print(f"\nContent checks:")
for check, result in checks.items():
    print(f"  {'✓' if result else '✗'} {check}")

print(f"\nTechnical checks:")
print(f"  Images: {len(images)} embedded ({', '.join([n.split('/')[-1] for n in images])})")
print(f"  OMML equations: {omml}")
print(f"  Raw $ symbols: {raw_dollar}")
print(f"  File size: {os.path.getsize(path)/1024:.0f} KB")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables: {len(doc.tables)}")

all_pass = all(checks.values()) and raw_dollar == 0 and len(images) == 3
print(f"\n  OVERALL: {'ALL PASS ✓' if all_pass else 'SOME CHECKS FAILED'}")
