#!/usr/bin/env python3
"""
AutoMCM-Pro: Word to LaTeX Converter
=====================================
Converts polished Word (.docx) to LaTeX (.tex) + compiles to PDF.

Usage:
    python scripts/docx_to_latex.py <input.docx>
    python scripts/docx_to_latex.py <input.docx> --output-dir ./output
    python scripts/docx_to_latex.py <input.docx> --no-compile

Pipeline: Word draft -> Human polish -> docx_to_latex.py -> LaTeX PDF
"""

import argparse, json, os, re, shutil, subprocess, sys, zipfile
from pathlib import Path

try:
    import docx
except ImportError:
    print("[ERROR] python-docx required: pip install python-docx")
    sys.exit(1)

KNOWN_ENGINES = [
    r"C:\Users\CX\AppData\Local\Programs\MiKTeX\miktex\bin\x64\pdflatex.exe",
    "pdflatex", "xelatex",
]

DOLLAR = chr(36)  # $ (avoids escape issues in replacement strings)

GREEK_MAP = {
    'α':'\\alpha','β':'\\beta','γ':'\\gamma','δ':'\\delta',
    'Δ':'\\Delta','ε':'\\epsilon','θ':'\\theta','λ':'\\lambda',
    'μ':'\\mu','σ':'\\sigma','φ':'\\phi','ω':'\\omega',
    'π':'\\pi','ρ':'\\rho','τ':'\\tau','η':'\\eta',
    'Σ':'\\Sigma','Π':'\\Pi',
}

# ============================================================
# MATH CONVERSION
# ============================================================

def word_math_to_latex(s):
    for g, l in GREEK_MAP.items(): s = s.replace(g, l)
    s = s.replace('×', ' \\times '); s = s.replace('·', ' \\cdot ')
    s = s.replace('°', '^{\\circ}')
    s = s.replace('≤', ' \\le '); s = s.replace('≥', ' \\ge ')
    s = s.replace('≠', ' \\neq '); s = s.replace('≈', ' \\approx ')
    s = s.replace('±', ' \\pm '); s = s.replace('∞', '\\infty')
    s = s.replace('∂', '\\partial')
    for func in ['sin','cos','tan','log','exp','ln','max','min','det','lim']:
        s = re.sub(r'(?<![\\a-zA-Z])' + func + r'(?![a-zA-Z])', r'\\' + func, s)
    s = re.sub(r'([a-zA-Z])_([a-zA-Z0-9]+)', r'\1_{\2}', s)
    s = re.sub(r'\^(-?\d+)', r'^{\1}', s)
    s = s.replace('²', '^{2}'); s = s.replace('³', '^{3}')
    return re.sub(r'\s+', ' ', s).strip()

def is_formula(text):
    if len(text) > 100: return False
    if re.match(r'^[\u4e00-\u9fff]', text):
        if any(kw in text for kw in ['其中','采用','利用','设置','定义','表示','计算',
                                       '包含','根据','引入','构建','建立','公式','如下',
                                       '考虑','结构','选取','选择','基于','通过']):
            return False
    if len(re.findall(r'[\u4e00-\u9fff]', text)) > 15: return False
    if not re.search(r'[=≈≠<>≤≥]', text): return False
    if not re.search(r'[αβγδa-zA-Z]', text): return False
    return True

# ============================================================
# TEXT PROCESSING (your core bug fix lives here!)
# ============================================================

def fix_inline_math(text):
    # 0: Fix $\alpha$_s -> $\alpha_s$
    text = re.sub(r'(\$\\[a-z]+\$)_([a-zA-Z0-9]+)',
                  lambda m: m.group(1) + DOLLAR + '_{' + m.group(2) + '}' + DOLLAR, text)
    # 1: R² etc
    text = re.sub(r'(?<!\$)([A-Za-z]+)\u00b2(?!\$)', lambda m: DOLLAR + m.group(1) + '^2' + DOLLAR, text)
    # 2: Greek letters
    for g, l in GREEK_MAP.items(): text = text.replace(g, DOLLAR + l + DOLLAR)
    # 3: Multi-char acronym subscripts: NWP_{DS} -> $NWP_{DS}$
    text = re.sub(r'(?<!\$)([A-Z]{2,})_(\{[^}]+\})',
                  lambda m: DOLLAR + m.group(1) + '_{' + m.group(2).strip('{}') + '}' + DOLLAR, text)
    text = re.sub(r'(?<!\$)([A-Z]{2,})_([a-zA-Z0-9]+)',
                  lambda m: DOLLAR + m.group(1) + '_{' + m.group(2) + '}' + DOLLAR, text)
    # 4: Chained subscripts: $NWP_{DS}$_GHI -> $NWP_{DS}$\_GHI
    text = re.sub(r'(\$[A-Z]+_\{[^}]+\}\$)_([a-zA-Z0-9]+)',
                  lambda m: m.group(1) + '\\_' + m.group(2), text)
    # 5: Single-char subscripts
    def sub_replacer(m):
        base, sub = m.group(1), m.group(2)
        if len(sub) <= 10 and re.match(r'^[a-zA-Z0-9]+\*?$', sub):
            return DOLLAR + base + '_{' + sub + '}' + DOLLAR
        return m.group(0)
    text = re.sub(r'(?<!\$)([a-zA-Z])_([a-zA-Z0-9]+)', sub_replacer, text)
    return text

def escape_text(text):
    parts = re.split(r'(\$[^$]+\$)', text)
    result = []
    for part in parts:
        if part.startswith('$') and part.endswith('$'):
            result.append(part)
        else:
            for ch, esc in [('&','\\&'),('%','\\%'),('#','\\#'),
                           ('{','\\{'),('}','\\}'),('~','\\textasciitilde ')]:
                part = part.replace(ch, esc)
            result.append(part)
    return ''.join(result)

def process_text(text):
    return escape_text(fix_inline_math(text))

# ============================================================
# TABLE BUILDING
# ============================================================

def build_table(tbl, caption):
    """Build a LaTeX longtable with \\small, balanced column widths."""
    if not tbl or not tbl[0]: return ""
    ncols = max(len(row) for row in tbl)
    
    # Calculate column widths based on content length
    col_max_lens = [0] * ncols
    for row in tbl:
        for ci in range(min(ncols, len(row))):
            col_max_lens[ci] = max(col_max_lens[ci], len(row[ci]))
    # Normalize to proportions
    total = sum(col_max_lens) or 1
    props = [max(0.08, cl / total) for cl in col_max_lens]
    # Renormalize
    total_p = sum(props)
    props = [p / total_p for p in props]
    
    # Build column spec
    col_specs = []
    for p in props:
        col_specs.append(f"p{{(\\linewidth - {ncols}\\tabcolsep) * \\real{{{p:.4f}}}}}")
    
    lines = ["{\\small\\begin{longtable}[]{@{}" + "".join(col_specs) + "@{}}"]
    lines.append(f"\\caption{{{escape_text(caption)}}}\\\\")
    lines.append("\\toprule\\noalign{}")
    
    # Header row
    header_cells = []
    for ci in range(ncols):
        c = process_text(tbl[0][ci]) if ci < len(tbl[0]) else ""
        if c.startswith("["): c = "{" + c + "}"
        header_cells.append(c)
    lines.append(" & ".join(header_cells) + " \\\\")
    lines.append("\\midrule\\noalign{}")
    lines.append("\\endhead")
    lines.append("\\bottomrule\\noalign{}")
    lines.append("\\endlastfoot")
    
    # Data rows
    for ri in range(1, len(tbl)):
        cells = []
        for ci in range(ncols):
            c = process_text(tbl[ri][ci]) if ci < len(tbl[ri]) else ""
            if c.startswith("["): c = "{" + c + "}"
            cells.append(c)
        lines.append(" & ".join(cells) + " \\\\")
    
    lines.append("\\end{longtable}}")
    return "\n".join(lines)
def extract_images(docx_path, output_dir, force=False):
    """Extract images from DOCX to figures/ subfolder.
    
    Only extracts if images don't already exist, avoiding duplicates.
    Set force=True to re-extract even if images exist.
    """
    img_dir = os.path.join(output_dir, "figures")
    os.makedirs(img_dir, exist_ok=True)
    
    # Check if images already exist (avoid re-extracting every time)
    existing = sorted([f for f in os.listdir(img_dir) if f.startswith("image") and f.endswith((".png",".jpg",".jpeg"))])
    if existing and not force:
        print(f"  [SKIP] {len(existing)} images already in figures/ -- use --force to re-extract")
        return existing
    
    images = []
    with zipfile.ZipFile(docx_path) as z:
        for name in sorted(z.namelist()):
            if name.startswith('word/media/'):
                data = z.read(name)
                ext = os.path.splitext(name)[1]
                img_name = f"image{len(images)+1}{ext}"
                with open(os.path.join(img_dir, img_name), 'wb') as f:
                    f.write(data)
                images.append(img_name)
    print(f"  [OK] Extracted {len(images)} images to figures/")
    return images

# ============================================================
# LATEX PREAMBLE
# ============================================================

LATEX_PREAMBLE = r'''\documentclass[12pt,a4paper]{ctexart}

% === Packages ===
\usepackage{amsmath,amssymb,amsfonts}
\usepackage{graphicx}
\usepackage{geometry}
\geometry{left=2.5cm,right=2.5cm,top=2.5cm,bottom=2.5cm}
\usepackage{booktabs,array,multirow,longtable,calc,ragged2e}
\usepackage{float}
\usepackage{enumitem}
\usepackage{fancyhdr}
\usepackage[font=bf,labelsep=quad,justification=centering]{caption}
\usepackage[colorlinks=false,hidelinks]{hyperref}
\usepackage{xcolor}
\usepackage{listings}
\usepackage{setspace}
\providecommand{\arraybackslash}{\textbackslash}

% === Page Setup ===
\setlength{\parindent}{2em}
\setlength{\parskip}{0.1em}
\setstretch{1.2}
\setlength{\headheight}{14pt}

% === Caption Setup ===
\captionsetup[table]{position=above,skip=6pt}
\captionsetup[longtable]{position=above,skip=6pt}
\captionsetup[figure]{position=below,skip=6pt}

% === Header/Footer ===
\pagestyle{fancy}
\fancyhf{}
\fancyhead[C]{\small 数学建模竞赛论文}
\fancyfoot[C]{\thepage}
\renewcommand{\headrulewidth}{0.4pt}

% === Code Style ===
\lstset{
    basicstyle=\ttfamily\small,
    breaklines=true,
    frame=single,
    numbers=left,
    numberstyle=\\tiny,
    xleftmargin=2em,
}

\\numberwithin{equation}{section}

\\begin{document}

'''

# ============================================================
# MAIN CONVERTER
# ============================================================

# ============================================================
# POST-PROCESSING
# ============================================================

def post_process_latex(latex):
    """Fix common artifacts: strip manual numbering, clean tables, fix refs, balance columns, add \\small."""
    
    # 1. Strip manual "图X "/"表X " from captions
    latex = re.sub(r'\\caption\{图\s*\d+\s+', r'\\caption{', latex)
    latex = re.sub(r'\\caption\{表\s*\d+\s+', r'\\caption{', latex)
    
    # 2. Replace in-text "图N" with \ref (only in document body)
    doc_start = latex.find(r'\begin{document}')
    if doc_start > 0:
        body = latex[doc_start:]
        body = re.sub(
            r'(?<!\\caption\{)(?<!\\label\{)图(\d+)(?=[^\d])',
            lambda m: r'图\\ref{fig:img' + m.group(1) + '}',
            body
        )
        latex = latex[:doc_start] + body
    
    # 3. Clean Pandoc table column spec artifacts
    latex = re.sub(r'>\{[^}]*\}', '', latex)
    latex = re.sub(r'\\arraybackslash', '', latex)
    
    # 4. Remove orphan \end{figure}
    lines = latex.split('\n')
    clean, depth = [], 0
    for line in lines:
        if r'\begin{figure}' in line: depth += 1
        if r'\end{figure}' in line:
            if depth > 0: depth -= 1
            else: continue
        clean.append(line)
    latex = '\n'.join(clean)
    
    # 5. Convert \[ \] to numbered equations
    latex = latex.replace(r'\[', r'\begin{equation}')
    latex = latex.replace(r'\]', r'\end{equation}')
    
    # 6. Remove duplicate equation wrappers
    latex = re.sub(r'\\begin\{equation\}\s*\\begin\{equation\}', r'\\begin{equation}', latex)
    latex = re.sub(r'\\end\{equation\}\s*\\end\{equation\}', r'\\end{equation}', latex)
    
    # 7. Remove empty captions
    latex = re.sub(r'\\caption\{\}\s*\n', '', latex)
    
    # 8. Wrap bare longtable (without \small) in \small
    latex = re.sub(
        r'(?<!\{\\small)(\\begin\{longtable\})',
        r'{\\small\1',
        latex
    )
    latex = re.sub(
        r'(\\end\{longtable\})(?!\})',
        r'\1}',
        latex
    )
    
    # 9. Ensure table/figure position specifiers
    latex = re.sub(r'\\begin\{figure\}(?!\[)', r'\\begin{figure}[H]', latex)
    
    # 10. Fix headheight warning if not already set
    if r'\setlength{\headheight}' not in latex:
        latex = latex.replace(
            r'\setstretch{1.2}',
            r'\setstretch{1.2}\n\setlength{\headheight}{14pt}'
        )
    
    return latex
def convert_docx_to_latex(docx_path, output_dir):
    doc = docx.Document(docx_path)
    # Extract images (skips if figures/ already has them)
    images = extract_images(docx_path, output_dir)

    paras = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        has_img = False
        for run in para.runs:
            drawings = run._element.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            if drawings: has_img = True
        if text or has_img:
            paras.append({"idx": i, "text": text, "has_image": has_img})

    tables = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(rows)

    latex = LATEX_PREAMBLE
    in_list = False
    img_idx = 0

    for p in paras:
        text = p["text"]
        if not text and not p["has_image"]: continue

        if p["idx"] == 0:
            latex += '\\begin{center}\n{\\LARGE\\heiti ' + text + '}\\\\[8pt]\n'
            continue
        if p["idx"] == 1 and not text.startswith("摘要") and not text.startswith("关键词"):
            latex += '{\\large ' + text + '}\n\\end{center}\n\\thispagestyle{empty}\n\n'
            continue

        if text == "摘要":
            latex += '\\begin{center}\n\\heiti\\Large 摘\\quad 要\n\\end{center}\n\n'
            continue
        if text.startswith("关键词"):
            kw = text.replace("关键词：", "").replace("关键词:", "")
            latex += '\\heiti{关键词：}' + kw + '\n\n'
            latex += '\\newpage\n\\tableofcontents\n\\newpage\n\n'
            continue

        sec_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
        if sec_match and len(text) < 40:
            num, title = sec_match.group(1), sec_match.group(2)
            dots = num.count('.')
            if dots == 0: latex += f'\\section{{{title}}}\n'
            elif dots == 1: latex += f'\\subsection{{{title}}}\n'
            else: latex += f'\\subsubsection{{{title}}}\n'
            continue

        table_cap = re.match(r'^表\s*(\d+)\s*(.*)$', text)
        if table_cap and len(text) < 40:
            cap_text = table_cap.group(2).strip() or f"表{int(table_cap.group(1))}"
            tbl_i = int(table_cap.group(1)) - 1
            if tbl_i < len(tables):
                latex += build_table(tables[tbl_i], cap_text) + '\n'
            continue

        list_match = re.match(r'^（(\d+)）(.+)', text)
        if list_match:
            if not in_list:
                latex += '\\begin{enumerate}[label=（\\arabic*）]\n'
                in_list = True
            latex += '\\item ' + process_text(list_match.group(2).strip()) + '\n'
            continue
        elif in_list:
            latex += '\\end{enumerate}\n\n'
            in_list = False

        if is_formula(text):
            fml = word_math_to_latex(text)
            lm = re.match(r'^(.+?)\s*[（(]\s*(\d+[a-z]?)\s*[）)]\s*$', fml)
            if lm:
                latex += '\\begin{equation}\n' + lm.group(1).strip() + \
                         '\n\\label{eq:' + lm.group(2) + '}\n\\end{equation}\n\n'
            else:
                latex += '\\begin{equation}\n' + fml + '\n\\end{equation}\n\n'
            continue

        if p["has_image"] and img_idx < len(images):
            latex += '\\begin{figure}[H]\n\\centering\n'
            latex += f'\\includegraphics[width=0.85\\textwidth]{{figures/{images[img_idx]}}}\n'
            if text:
                cap = re.sub(r'^图\s*\d+\s*', '', text)
                latex += f'\\caption{{{escape_text(cap)}}}\n'
            latex += '\\end{figure}\n\n'
            img_idx += 1
            continue

        if text:
            latex += process_text(text) + '\n\n'

    if in_list: latex += '\\end{enumerate}\n\n'
    latex += '\\end{document}\n'

    tex_name = os.path.splitext(os.path.basename(docx_path))[0] + ".tex"
    tex_path = os.path.join(output_dir, tex_name)
    # Clean Pandoc artifacts
        # Remove orphan \end{figure}
    import re as _re
    latex = _re.sub(r'(?<!\\begin\{figure\}.*?)\\end\{figure\}', '', latex, flags=_re.DOTALL)
    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(latex)
    print(f"[convert] LaTeX: {tex_path} ({len(latex)} bytes)")
    return tex_path

# ============================================================
# PDF COMPILATION
# ============================================================

def find_engine():
    for eng in KNOWN_ENGINES:
        if Path(eng).exists() or shutil.which(eng):
            return eng
    return None

def compile_tex(tex_path):
    engine = find_engine()
    if not engine:
        print("[WARNING] No LaTeX engine found. Skipping PDF compilation.")
        return None
    tex_dir = os.path.dirname(tex_path) or "."
    tex_file = os.path.basename(tex_path)
    cmd = [engine, "-synctex=1", "-interaction=nonstopmode", "-file-line-error", tex_file]
    print(f"[compile] Engine: {engine}")
    for pn in [1, 2]:
        print(f"[compile] Pass {pn}...")
        result = subprocess.run(cmd, cwd=tex_dir, capture_output=True,
                               text=True, encoding="utf-8", errors="replace")
        for line in (result.stdout + result.stderr).splitlines():
            if any(k in line for k in ("Error", "Fatal", "Output written", "Warning:")):
                if "major issue" not in line: print(f"  {line}")
    pdf_path = tex_path.replace(".tex", ".pdf")
    if os.path.exists(pdf_path):
        size_kb = os.path.getsize(pdf_path) / 1024
        print(f"[compile] SUCCESS: {pdf_path} ({size_kb:.0f} KB)")
        return pdf_path
    else:
        print("[compile] FAILED - check .log file")
        return None

# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="AutoMCM-Pro: Word -> LaTeX PDF")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("--output-dir", "-o", default=None, help="Output directory")
    parser.add_argument("--no-compile", action="store_true", help="Skip PDF compilation")
    parser.add_argument("--cleanup", action="store_true", help="Auto-remove intermediate files (.tex, .aux, .log)")
    args = parser.parse_args()

    docx_path = os.path.abspath(args.input)
    if not os.path.exists(docx_path):
        sys.exit(f"[ERROR] File not found: {docx_path}")
    output_dir = args.output_dir or os.path.dirname(docx_path)
    os.makedirs(output_dir, exist_ok=True)

    print(f"[convert] Input:  {docx_path}")
    tex_path = convert_docx_to_latex(docx_path, output_dir)
    if not args.no_compile:
        pdf_path = compile_tex(tex_path)
        if pdf_path:
            print(f"\nDone! PDF: {pdf_path}")
            if args.cleanup:
                # Clean up intermediate files
                tex_dir = os.path.dirname(tex_path) or "."
                tex_name = os.path.splitext(os.path.basename(tex_path))[0]
                for ext in [".aux", ".log", ".toc", ".synctex.gz", ".out"]:
                    artifact = os.path.join(tex_dir, tex_name + ext)
                    if os.path.exists(artifact):
                        try: os.remove(artifact)
                        except: pass
                # Remove .tex itself if cleanup requested
                try: os.remove(tex_path)
                except: pass
                print("[cleanup] Removed intermediate files")
        else:
            print(f"\nTeX ready for manual compile: {tex_path}")
    else:
        print(f"\nTeX ready: {tex_path}")

if __name__ == "__main__":
    main()

