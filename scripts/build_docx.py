#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown to DOCX builder with OMML equations, three-line tables with captions."""
import sys
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')
import argparse, re, os
from pathlib import Path

# Dependencies
for pkg, mod in [("python-docx","docx"),("latex2mathml","latex2mathml.converter"),("mathml2omml","mathml2omml"),("lxml","lxml")]:
    try: __import__(mod)
    except ImportError:
        print(f"[ERROR] Missing: {pkg}"); sys.exit(1)

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from latex2mathml.converter import convert as l2m_convert
from mathml2omml import convert as m2o
from lxml import etree

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# ── OMML Pipeline ────────────────────────────────────────────
def latex_to_omml(latex):
    latex = latex.strip()
    if not latex: return ""
    try: return m2o(l2m_convert(latex))
    except Exception as e:
        print(f"  [WARN] LaTeX->OMML: {latex[:50]}... ({e})")
        return ""

def inject_omml(para, omml_xml):
    if not omml_xml: return False
    try:
        omml_xml = omml_xml.replace("<m:oMath>", f'<m:oMath xmlns:m="{MATH_NS}">')
        para._p.append(etree.fromstring(omml_xml.encode("utf-8")))
        return True
    except Exception as e:
        print(f"  [WARN] OMML inject: {e}")
        return False

# ── Markdown cleanup ─────────────────────────────────────────
def clean_md_text(text):
    """Remove Markdown syntax that should not appear in the final DOCX text."""
    text = str(text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = text.replace("\\*", "*").replace("\\_", "_").replace("\\`", "`")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^\s*>\s*", "", text)
    text = re.sub(r"^\s*[-*+]\s+", "", text)
    text = re.sub(r"^\s*\d+[.)]\s+", "", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(?<!\w)(\*|_)([^*_]+)\1(?!\w)", r"\2", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def add_markdown_runs(paragraph, text, size=12):
    """Add cleaned text while preserving simple bold markdown."""
    text = str(text)
    pos = 0
    pattern = re.compile(r"(\*\*|__)(.+?)\1")
    for m in pattern.finditer(text):
        before = clean_md_text(text[pos:m.start()])
        if before:
            r = paragraph.add_run(before)
            set_font(r, size)
        bold_text = clean_md_text(m.group(2))
        if bold_text:
            r = paragraph.add_run(bold_text)
            set_font(r, size, bold=True)
        pos = m.end()
    rest = clean_md_text(text[pos:])
    if rest:
        r = paragraph.add_run(rest)
        set_font(r, size)

def count_substantive_chars(text):
    """Approximate paper substance, excluding Markdown/code scaffolding."""
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"\$\$.*?\$\$", "", text, flags=re.DOTALL)
    text = re.sub(r"\|[-:\s|]+\|?", "", text)
    text = clean_md_text(text)
    return len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text))

# ── Three-line Table ─────────────────────────────────────────
def _add_table_cell_text(cell, text):
    """Add text to a table cell, converting $...$ inline formulas to OMML."""
    INLINE_EQ = re.compile(r"\$(.+?)\$")
    segments = INLINE_EQ.split(str(text))
    
    # Clear cell
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for k, seg in enumerate(segments):
        if k % 2 == 1:
            # Equation segment - convert to OMML
            omml = latex_to_omml(seg)
            if omml:
                inject_omml(p, omml)
            else:
                r = p.add_run("$" + clean_md_text(seg) + "$")
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(9)
        else:
            # Text segment
            cleaned = clean_md_text(seg)
            if cleaned:
                r = p.add_run(cleaned)
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(9)

def add_three_line_table(doc, headers, rows, caption_text=None):
    """CUMCM three-line table with optional caption."""
    if caption_text:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(clean_md_text(caption_text))
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10)
        r.font.bold = True

    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    # Borders: top+bottom thick, no vertical
    tbl = table._tbl
    tblPr = tbl.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr")
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for tag, val, sz in [("top","single","12"),("bottom","single","12"),
                          ("left","none","0"),("right","none","0"),
                          ("insideH","none","0"),("insideV","none","0")]:
        e = OxmlElement(f"w:{tag}"); e.set(qn("w:val"),val); e.set(qn("w:sz"),sz)
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),"000000"); borders.append(e)
    tblPr.append(borders)

    # Header row - with inline formula support
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _add_table_cell_text(cell, h)
        # Make header bold
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
        # Thin bottom border on header cells
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        be = OxmlElement("w:bottom"); be.set(qn("w:val"),"single")
        be.set(qn("w:sz"),"4"); be.set(qn("w:color"),"000000")
        tcBorders.append(be); tcPr.append(tcBorders)

    # Data rows - with inline formula support
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            _add_table_cell_text(cell, str(val))

    doc.add_paragraph()  # spacing after table
    return table

# ── Markdown Table Parser ────────────────────────────────────
def parse_md_table(lines, idx):
    """Parse Markdown table. Returns (headers, rows, next_idx) or (None,None,idx)."""
    if idx >= len(lines): return None, None, idx
    hl = lines[idx].strip()
    if "|" not in hl: return None, None, idx
    headers = [h.strip() for h in hl.split("|") if h.strip()]
    if idx+1 >= len(lines): return None, None, idx
    sep = lines[idx+1].strip()
    if not re.match(r"^[\|\s\-:]+$", sep): return None, None, idx

    rows = []
    j = idx + 2
    while j < len(lines):
        rl = lines[j].strip()
        if "|" not in rl: break
        vals = [v.strip() for v in rl.split("|") if v.strip()]
        if len(vals) >= len(headers):
            rows.append(vals[:len(headers)])
        elif len(vals) > 0:
            rows.append(vals + [""]*(len(headers)-len(vals)))
        j += 1
    return headers, rows, j

# ── Font helper ──────────────────────────────────────────────
def set_font(run, size=12, bold=False, cn="宋体", en="Times New Roman"):
    run.font.name = en; run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size); run.font.bold = bold

# ── Main Builder ─────────────────────────────────────────────
def build_paper(md_path, output_path):
    with open(md_path, "r", encoding="utf-8") as f:
        raw = f.read()
    lines = raw.split("\n")
    source_chars = count_substantive_chars(raw)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.18); s.right_margin = Cm(3.18)

    i = 0; eq_ok = eq_fail = tbl_n = fig_n = 0
    INLINE_EQ = re.compile(r"\$(.+?)\$")
    in_code_block = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Code fence - skip fenced code content entirely.
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # Skip empty
        if not line: i += 1; continue

        # Horizontal rule / separator - skip silently
        if line.strip() in ("---", "***", "___"):
            i += 1; continue

        # Heading
        if line.startswith("## "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
            r = p.add_run(clean_md_text(line[3:])); set_font(r, 14, True, "黑体"); i += 1; continue
        if line.startswith("### "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
            r = p.add_run(clean_md_text(line[4:])); set_font(r, 12, True, "黑体"); i += 1; continue
        if line.startswith("#### "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
            r = p.add_run(clean_md_text(line[5:])); set_font(r, 12, True, "宋体"); i += 1; continue
        if line.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(40)
            r = p.add_run(clean_md_text(line[2:])); set_font(r, 18, True, "黑体"); i += 1; continue

        # Table
        if "|" in line and i+1 < len(lines) and re.match(r"^[\|\s\-:]+$", lines[i+1].strip()):
            # Check for caption in preceding non-empty line
            caption = None
            j = i - 1
            while j >= 0 and not lines[j].strip(): j -= 1
            if j >= 0:
                prev = lines[j].strip()
                cm = re.match(r"(?:\*\*)?表\s*(\d+)\s*[：:\s]?\s*(.+?)(?:\*\*)?$", prev)
                if cm:
                    caption = f"表{cm.group(1)} {clean_md_text(cm.group(2))}"
                    # Remove caption line from output (already consumed)
                    lines[j] = ""

            headers, rows, next_i = parse_md_table(lines, i)
            if headers and rows:
                tbl_n += 1
                if caption is None:
                    caption = f"表{tbl_n}"
                add_three_line_table(doc, headers, rows, caption)
                i = next_i; continue

        # Display equation: $$...$$ (single or multi-line)
        if line.startswith("$$"):
            if len(line) > 4 and line.endswith("$$"):
                # Single-line $$...$$
                latex = line[2:-2].strip()
            else:
                # Multi-line $$ block
                latex_lines = []
                j = i + 1
                while j < len(lines) and not lines[j].strip() == "$$":
                    latex_lines.append(lines[j])
                    j += 1
                latex = "\n".join(latex_lines).strip()
                i = j  # skip to closing $$
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            omml = latex_to_omml(latex)
            if omml and inject_omml(p, omml):
                eq_ok += 1; print(f"  [EQ-{eq_ok}] {latex[:60]}...")
            else:
                eq_fail += 1; r = p.add_run(f"[公式: {latex[:40]}...]"); set_font(r, 10)
            i += 1; continue

        # FIGURE: [FIGURE: path | caption]
        fm = re.match(r"\[FIGURE:\s*(.+?)\s*\|\s*(.+?)\]", line)
        if fm:
            fname = fm.group(1).strip(); cap_text = clean_md_text(fm.group(2))
            if os.path.exists(fname):
                fig_n += 1
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    r = p.add_run(); r.add_picture(fname, width=Inches(5.0))
                    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = cap.add_run(cap_text); set_font(rc, 10)
                    print(f"  [FIG-{fig_n}] {os.path.basename(fname)}")
                except Exception as e:
                    print(f"  [WARN] Image failed: {fname} - {e}")
            else:
                print(f"  [WARN] Image not found: {fname}")
            i += 1; continue

        # Regular paragraph - process **bold** and $inline$ equations
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.5

        # First, process inline equations: split by $...$
        eq_segments = INLINE_EQ.split(line)
        for k, seg in enumerate(eq_segments):
            if k % 2 == 1:
                # Equation segment
                omml = latex_to_omml(seg)
                if omml:
                    inject_omml(p, omml); eq_ok += 1
                else:
                    r = p.add_run(f"${seg}$"); set_font(r, 10)
            else:
                add_markdown_runs(p, seg, 12)
        i += 1

    # Content length check
    all_text = " ".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    chars = len(all_text)
    HARD_MIN_CHARS = 9000
    print(f"\n[CHECK] Total chars: {chars} | source substance: {source_chars}")
    print(f"        hard minimum {HARD_MIN_CHARS}; no artificial upper/target length")
    if chars < HARD_MIN_CHARS:
        print(f"  HARD FAIL: {chars} < {HARD_MIN_CHARS}. The draft is too thin to submit.")
        print("  Expand with derivations, baseline comparison, robustness, figure interpretation, and error analysis.")
        return None
    print("  PASS: content clears the minimum. Do not pad; let the problem depth determine final length.")

    doc.save(output_path)
    pages = chars / 900
    print(f"[DONE] {output_path}")
    print(f"  Equations: {eq_ok} OK, {eq_fail} FAIL")
    print(f"  Tables: {tbl_n} | Figures: {fig_n}")
    print(f"  Total chars: {chars} | Pages: ~{pages:.0f}")
    return output_path

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("markdown"); ap.add_argument("output")
    args = ap.parse_args()
    build_paper(args.markdown, args.output)


