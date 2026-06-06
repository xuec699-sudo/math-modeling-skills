#!/usr/bin/env python3
"""
Math Modeling Contest - Quality Gate
Five-layer quality check before advancing any stage.

Exit codes:
  0 = All passed, can advance
  1 = Gate failed, cannot advance
  2 = Gate skipped (not applicable to current stage)

Usage (by agent, not user-facing):
  python scripts/quality_gate.py verify   --stage model_1_verify --report-file REPORT
  python scripts/quality_gate.py sanity   --stage model_1_build  --output-file OUTPUT
  python scripts/quality_gate.py lit      --stage model_1_build  --problem-n 1
  python scripts/quality_gate.py consist  --problems 3
  python scripts/quality_gate.py all      --stage model_1_verify --report-file REPORT
"""

import argparse
import re
import sys
from pathlib import Path

WORKSPACE      = Path("CUMCM_Workspace")
THOUGHT_FILE   = WORKSPACE / "memory" / "thought_process.md"

# KyrieZhang329-inspired gate contracts (v5.7.0)
try:
    from gate_contracts import (
        GATE_CONTRACTS, print_gate_contracts,
        propagation_check, print_propagation_report,
        gate_g2_poc, gate_g4_frozen_staleness,
        audit_consistency_enhanced, audit_completeness_enhanced, audit_quality_enhanced,
        run_g6_audit_enhanced, print_g6_enhanced_report,
    )
    _HAS_GATE_CONTRACTS = True
except ImportError:
    _HAS_GATE_CONTRACTS = False



# ---------------------------------------------------------------------------
# Gate 6: Render QA Gate
# ---------------------------------------------------------------------------

def gate_render_qa(render_dir: str = None) -> tuple[bool, str]:
    """Check that DOCX render PNGs exist and are valid for visual inspection."""
    if render_dir is None:
        render_dir = WORKSPACE / "output" / "qa"
    rd = Path(render_dir)
    
    if not rd.exists():
        return False, (
            f"Render QA directory not found: {rd}\n"
            "Run: python render_docx.py output/<paper>.docx --output_dir output/qa/"
        )
    
    pngs = sorted(rd.glob("page-*.png"))
    if not pngs:
        return False, f"No page-*.png files found in {rd}. Rendering may have failed."
    
    bad = [p.name for p in pngs if p.stat().st_size == 0]
    if bad:
        return False, f"Zero-byte PNGs detected: {bad}. Re-run render_docx.py."
    
    page_count = len(pngs)
    if page_count > 30:
        return False, f"Page count {page_count} exceeds competition limit (30 pages)."
    
    return True, (
        f"Render artifacts OK ({page_count} pages).\n"
        f"VISUAL INSPECTION REQUIRED: Open {rd}/page-*.png at 100% zoom and verify:\n"
        f"  - All fonts correct (SimHei/Song/TNR)\n"
        f"  - No clipped/overlapping text\n"
        f"  - Three-line table borders correct\n"
        f"  - Equations rendered (OMML not broken)\n"
        f"  - Images clear and correctly placed\n"
        f"  - Page order correct, no missing sections\n"
        f"Only advance to [final_deliver] after ALL pages pass visual check."
    )# ---------------------------------------------------------------------------
# Gate 1: Verification Report Parsing
# ---------------------------------------------------------------------------

def gate_verify_report(report_file: str) -> tuple[bool, str]:
    """Parse structured VERIFICATION REPORT from verify_*.py output."""
    path = Path(report_file)
    if not path.exists():
        return False, f"Verification report file not found: {path}"

    text = path.read_text(encoding="utf-8", errors="replace")

    report_block = re.search(
        r"={5,}\s*VERIFICATION REPORT\s*={5,}(.*?)={5,}",
        text, re.DOTALL | re.IGNORECASE
    )
    if not report_block:
        return False, (
            "Verification report format invalid: missing '===VERIFICATION REPORT===' block.\n"
            "verify_*.py must print a structured report at the end."
        )

    block = report_block.group(1)

    result_match = re.search(r"Result\s*:\s*(PASS|FAIL)", block, re.IGNORECASE)
    if not result_match:
        return False, "Report missing 'Result: PASS' or 'Result: FAIL' line."

    result = result_match.group(1).upper()
    if result == "FAIL":
        fail_items = re.findall(r"\u2717\s+(.+)", block)
        detail = "\n  ".join(fail_items) if fail_items else "(see report details)"
        return False, f"Verification FAILED. Failed items:\n  {detail}"

    return True, "Verification PASSED \u2714"


# ---------------------------------------------------------------------------
# Gate 2: Numerical Sanity Check
# ---------------------------------------------------------------------------

_BAD_PATTERNS = [
    (re.compile(r'\binf\b', re.IGNORECASE),   "Output contains inf"),
    (re.compile(r'\bnan\b', re.IGNORECASE),   "Output contains nan"),
    (re.compile(r'1\.?\d*[eE][+]?[23][0-9]{2,}'), "Extreme magnitude (>=1e200)"),
    (re.compile(r'-1\.?\d*[eE][+]?[23][0-9]{2,}'), "Extreme negative magnitude (<=-1e200)"),
]


def gate_numerical_sanity(output_file: str) -> tuple[bool, str]:
    """Scan model output for inf/nan and extreme values."""
    path = Path(output_file)
    if not path.exists():
        return True, "Output file not found, skipping numerical check (verify manually)"

    text = path.read_text(encoding="utf-8", errors="replace")
    issues = []
    for pat, desc in _BAD_PATTERNS:
        matches = pat.findall(text)
        if matches:
            issues.append(f"{desc} (found {len(matches)} occurrences)")

    if issues:
        return False, "Numerical sanity FAILED:\n  " + "\n  ".join(issues)
    return True, "Numerical sanity PASSED \u2714"


# ---------------------------------------------------------------------------
# Gate 3: Literature Citation Count
# ---------------------------------------------------------------------------

def gate_literature(problem_n: int) -> tuple[bool, str]:
    """Check that at least 2 academic references support each sub-problem's model choice."""
    if not THOUGHT_FILE.exists():
        return False, "thought_process.md not found - cannot verify literature"

    text = THOUGHT_FILE.read_text(encoding="utf-8", errors="replace")

    # Look for model justification for problem N with references
    section_pattern = re.compile(
        rf"(?:Problem|问题|Q)\s*{problem_n}[^#]*?(?=(?:Problem|问题|Q)\s*{problem_n + 1}|\Z)",
        re.DOTALL | re.IGNORECASE
    )
    section_match = section_pattern.search(text)
    if not section_match:
        # Try broader search
        refs = re.findall(r'(?:doi|DOI|arXiv|ISBN|http)[^\s]{5,}', text)
        if len(refs) >= 2 * problem_n:
            return True, f"Literature PASSED \u2714 (found {len(refs)} references total)"
        return False, f"Cannot verify literature for problem {problem_n}. Ensure at least 2 references per sub-problem."

    section = section_match.group(0)
    refs_in_section = re.findall(r'(?:doi|DOI|arXiv|ISBN|http|\\cite|\\ref|Ref\.|Reference)', section)

    if len(refs_in_section) >= 2:
        return True, f"Literature PASSED \u2714 (problem {problem_n}: {len(refs_in_section)} refs)"

    return False, (
        f"Literature FAILED for problem {problem_n}: only {len(refs_in_section)} references found.\n"
        f"Requirement: at least 2 academic references per sub-problem supporting model choice."
    )


# ---------------------------------------------------------------------------
# Gate 4: Multi-Problem Consistency
# ---------------------------------------------------------------------------

def gate_consistency(num_problems: int) -> tuple[bool, str]:
    """Check that physical constants are consistent across sub-problems."""
    import ast

    models_dir = WORKSPACE / "src" / "models"
    if not models_dir.exists():
        return True, "Models directory not found, skipping consistency check"

    py_files = sorted(models_dir.glob("problem*.py"))
    if len(py_files) < 2:
        return True, "Only one model file, skipping consistency check \u2714"

    # Scan for CONSTANT = value patterns
    const_pattern = re.compile(r'^([A-Z][A-Z_0-9]{2,})\s*=\s*([0-9.eE+\-]+)', re.MULTILINE)
    const_map: dict[str, dict[str, str]] = {}

    for py_file in py_files:
        try:
            text = py_file.read_text(encoding="utf-8")
            for match in const_pattern.finditer(text):
                name = match.group(1)
                value = match.group(2).strip()
                const_map.setdefault(name, {})[py_file.name] = value
        except Exception:
            continue

    conflicts = []
    for name, file_vals in const_map.items():
        unique_vals = set(file_vals.values())
        if len(unique_vals) > 1:
            detail = ", ".join(f"{f}={v}" for f, v in file_vals.items())
            conflicts.append(f"  {name}: {detail}")

    if conflicts:
        return False, (
            "Cross-problem consistency FAILED - differing physical constants:\n"
            + "\n".join(conflicts)
            + "\nPlease unify and re-run."
        )
    return True, f"Consistency PASSED \u2714 (scanned {len(py_files)} files, {len(const_map)} constants)"


# ---------------------------------------------------------------------------
# Gate 5: Model-Vefiry Pair Check
# ---------------------------------------------------------------------------

def gate_verify_pairing(stage: str) -> tuple[bool, str]:
    """Ensure every model_N_build has a corresponding verify_N script."""
    models_dir = WORKSPACE / "src" / "models"
    verify_dir = WORKSPACE / "src" / "verifications"

    if not models_dir.exists():
        return True, "Models directory not found, skipping pairing check"

    model_files = sorted(models_dir.glob("problem*.py"))
    missing_verifies = []

    for mf in model_files:
        # Expected: problem1_optimization.py -> verify_problem1_optimization.py
        stem = mf.stem  # e.g., problem1_optimization
        verify_name = f"verify_{stem}.py"
        if not (verify_dir / verify_name).exists():
            # Try broader match: verify_problem1_*.py
            broader = list(verify_dir.glob(f"verify_{stem.split('_')[0]}_*.py"))
            if not broader:
                missing_verifies.append(mf.name)

    if missing_verifies:
        return False, (
            f"Missing verification scripts for:\n  "
            + "\n  ".join(missing_verifies)
            + "\nEvery model must have a corresponding verify_*.py script."
        )

    return True, f"Model-verify pairing PASSED \u2714 ({len(model_files)} models, all paired)"


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

def run_gate(name: str, args) -> bool:
    """Run a single gate, print result, return pass/fail."""
    if name == "verify":
        passed, msg = gate_verify_report(args.report_file)
    elif name == "sanity":
        passed, msg = gate_numerical_sanity(args.output_file)
    elif name == "lit":
        passed, msg = gate_literature(args.problem_n)
    elif name == "consist":
        passed, msg = gate_consistency(args.problems)
    elif name == 'render':
        passed, msg = gate_render_qa(args.render_dir)
    elif name == 'pairing':
        passed, msg = gate_verify_pairing(args.stage)
    elif name == 'contracts':
        if _HAS_GATE_CONTRACTS:
            print_gate_contracts()
        return True
    elif name == 'g2_poc':
        if _HAS_GATE_CONTRACTS:
            passed, msg, _ = gate_g2_poc()
            prefix = "\u2714 GATE" if passed else "\u2717 GATE"
            print(f"{prefix} [g2_poc] {msg}")
            return passed
        return True
    elif name == 'g4_stale':
        if _HAS_GATE_CONTRACTS:
            subq = getattr(args, 'subquestion', None)
            passed, msg, _ = gate_g4_frozen_staleness(subq)
            prefix = "\u2714 GATE" if passed else "\u2717 GATE"
            print(f"{prefix} [g4_stale] {msg}")
            return passed
        return True
    elif name == 'g6_audit':
        if _HAS_GATE_CONTRACTS:
            ws = getattr(args, 'workspace_path', None)
            results = run_g6_audit_enhanced(ws)
            print_g6_enhanced_report(results)
            return results.get('gate_passed', False)
        return True
    elif name == 'propagate':
        if _HAS_GATE_CONTRACTS:
            changed = getattr(args, 'changed_files', [])
            if changed:
                result = propagation_check(changed if isinstance(changed, list) else [changed])
                print_propagation_report(result)
        return True
    else:
        return True

    prefix = "\u2714 GATE" if passed else "\u2717 GATE"
    print(f"{prefix} [{name}] {msg}")
    return passed



# ---------------------------------------------------------------------------
# Gate 7: Content Depth Check (NEW)
# ---------------------------------------------------------------------------

CONTENT_DEPTH_RULES = {
    # (section_keyword, min_chars, min_paragraphs)
    "problem": [
        ("问题重述", 300, 3),
        ("问题分析", 500, 5),
        ("数据预处理", 400, 4),
    ],
    "per_problem": [
        ("模型建立", 400, 4),
        ("求解", 400, 4),
        ("结果", 300, 3),
    ],
}


def gate_content_depth(docx_path: str) -> tuple[bool, str]:
    """Check paper DOCX has sufficient content depth per section."""
    try:
        from docx import Document
    except ImportError:
        return True, "[SKIP] python-docx not available, skipping content depth check"

    path = Path(docx_path)
    if not path.exists():
        return False, f"DOCX not found: {docx_path}"

    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    issues = []

    # Global checks
    for keyword, min_chars, min_paras in CONTENT_DEPTH_RULES["problem"]:
        idx = full_text.find(keyword)
        if idx == -1:
            issues.append(f"MISSING section: '{keyword}'")
            continue
        # Get section text (rough: next 2000 chars or until next major section)
        section_end = min(idx + 3000, len(full_text))
        section_text = full_text[idx:section_end]
        paras = [l for l in section_text.split("\n") if l.strip()]
        if len(section_text) < min_chars:
            issues.append(f"'{keyword}' too short: {len(section_text)} chars (min {min_chars})")
        if len(paras) < min_paras:
            issues.append(f"'{keyword}' too few paragraphs: {len(paras)} (min {min_paras})")

    # Per-problem checks: find sections labeled 四、五、六、七、八
    problem_markers = ["四、", "五、", "六、", "七、"]
    for marker in problem_markers:
        idx = full_text.find(marker)
        if idx == -1:
            continue
        # Find next marker or end
        next_idx = len(full_text)
        for m2 in problem_markers + ["八、", "九、", "参考文献"]:
            mi = full_text.find(m2, idx + 2)
            if mi != -1 and mi < next_idx:
                next_idx = mi
        section_text = full_text[idx:next_idx]
        paras = [l for l in section_text.split("\n") if l.strip()]
        if len(section_text) < 800:
            issues.append(f"{marker} section too short: {len(section_text)} chars (min 800)")
        if len(paras) < 8:
            issues.append(f"{marker} section too few paragraphs: {len(paras)} (min 8)")

    if issues:
        return False, "Content depth FAILED:\n  - " + "\n  - ".join(issues)

    return True, "Content depth PASSED - all sections have sufficient content"


# ---------------------------------------------------------------------------
# Gate 8: Formula Presence Check (NEW)
# ---------------------------------------------------------------------------

FORMULA_PATTERNS = [
    r'[=≈≠≤≥]',           # equation signs
    r'[αβγδεθλμπσφω]',     # Greek letters
    r'[∑∏∫∂√∞]',           # math operators
    r'R\s*[²2]',            # R-squared
    r'p\s*[<≤]',            # p-value
    r'[A-Z]\s*=',           # variable = ...
    r'\\frac', r'\\sum', r'\\int',  # LaTeX
    r'beta', r'sigma', r'alpha',    # English math terms
    r'logit|sigmoid|softmax|argmax|argmin',
    r'\d+\.\d+e[+-]\d+',    # scientific notation
]


def gate_formula_presence(docx_path: str) -> tuple[bool, str]:
    """Check each problem section has at least one mathematical formula/expression."""
    try:
        from docx import Document
    except ImportError:
        return True, "[SKIP] python-docx not available, skipping formula check"

    path = Path(docx_path)
    if not path.exists():
        return False, f"DOCX not found: {docx_path}"

    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    problem_markers = ["四、", "五、", "六、", "七、"]
    issues = []

    for marker in problem_markers:
        idx = full_text.find(marker)
        if idx == -1:
            continue
        next_idx = len(full_text)
        for m2 in problem_markers + ["八、", "九、", "参考文献"]:
            mi = full_text.find(m2, idx + 2)
            if mi != -1 and mi < next_idx:
                next_idx = mi
        section_text = full_text[idx:next_idx]

        found = False
        for pattern in FORMULA_PATTERNS:
            if re.search(pattern, section_text, re.IGNORECASE):
                found = True
                break

        if not found:
            issues.append(f"{marker} section has NO mathematical formulas/expressions")

    if issues:
        return False, "Formula presence FAILED:\n  - " + "\n  - ".join(issues)

    return True, "Formula presence PASSED - all problem sections contain mathematical expressions"


# ---------------------------------------------------------------------------
# Gate 9: Figure Context Check (NEW)
# ---------------------------------------------------------------------------

def gate_figure_context(docx_path: str) -> tuple[bool, str]:
    """Check every figure in the paper has explanatory text nearby."""
    try:
        from docx import Document
    except ImportError:
        return True, "[SKIP] python-docx not available, skipping figure context check"

    path = Path(docx_path)
    if not path.exists():
        return False, f"DOCX not found: {docx_path}"

    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    # Find figure references (图N or 图 N)
    fig_refs = re.findall(r'图\s*(\d+)', full_text)
    if not fig_refs:
        return True, "No figures found, skipping figure context check"

    issues = []
    for fig_num in set(fig_refs):
        pattern = rf'图\s*{fig_num}'
        matches = list(re.finditer(pattern, full_text))
        for m in matches:
            pos = m.end()
            # Check next 500 chars for explanatory text (at least 50 chars)
            after = full_text[pos:pos + 500]
            # Skip the caption line (usually short)
            lines = after.split('\n')
            explanation = ""
            for line in lines:
                if len(line.strip()) > 30:
                    explanation = line.strip()
                    break
            if len(explanation) < 30:
                issues.append(f"图{fig_num}: no explanatory text found after figure caption")

    if issues:
        return False, "Figure context FAILED:\n  - " + "\n  - ".join(issues)

    return True, f"Figure context PASSED - all {len(set(fig_refs))} figure(s) have explanatory text"


# ---------------------------------------------------------------------------
# Gate 10: Model Quality Threshold Check (NEW)
# ---------------------------------------------------------------------------

MODEL_QUALITY_RULES = {
    "regression": {"min_r2": 0.05, "min_pvalue": 0.05, "remediation": "Try nonlinear transforms (log, sqrt, Box-Cox) and interaction terms"},
    "classification": {"min_auc": 0.70, "min_recall": 0.20, "remediation": "Try SMOTE oversampling, class weighting, or feature engineering"},
    "clustering": {"min_silhouette": 0.30, "remediation": "Try different k values or alternative clustering algorithms"},
}


def gate_model_quality(report_text: str, model_type: str = "regression") -> tuple[bool, str]:
    """Check model results meet minimum quality thresholds."""
    rules = MODEL_QUALITY_RULES.get(model_type, {})
    issues = []

    if model_type == "regression":
        # Extract R-squared
        r2_match = re.search(r'R[²2]\s*[=:]\s*([0-9.]+)', report_text, re.IGNORECASE)
        adj_r2_match = re.search(r'Adj[^=]*R[²2]\s*[=:]\s*([0-9.]+)', report_text, re.IGNORECASE)
        cv_r2_match = re.search(r'CV\s*R[²2]\s*[=:]\s*([-0-9.]+)', report_text, re.IGNORECASE)

        if r2_match:
            r2 = float(r2_match.group(1))
            if r2 < rules.get("min_r2", 0.05):
                issues.append(f"R2={r2:.3f} below threshold {rules['min_r2']}. {rules.get('remediation', '')}")
            else:
                pass  # OK

        if cv_r2_match:
            cv_r2 = float(cv_r2_match.group(1))
            if cv_r2 < -0.5:
                issues.append(f"CV R2={cv_r2:.3f} is severely negative - model has no predictive power. Consider alternative approaches or acknowledge limitation.")

    elif model_type == "classification":
        auc_match = re.search(r'AUC\s*[=:]\s*([0-9.]+)', report_text, re.IGNORECASE)
        recall_match = re.search(r'recall\s*[=:]\s*([0-9.]+)', report_text, re.IGNORECASE)

        if auc_match and float(auc_match.group(1)) < rules.get("min_auc", 0.70):
            issues.append(f"AUC below threshold. {rules.get('remediation', '')}")
        if recall_match and float(recall_match.group(1)) < rules.get("min_recall", 0.20):
            issues.append(f"Recall too low. {rules.get('remediation', '')}")

    if issues:
        return False, "Model quality FAILED:\n  - " + "\n  - ".join(issues)

    return True, "Model quality PASSED - metrics within acceptable range"




# ---------------------------------------------------------------------------  
# Gate 11: Model Formulation Completeness
# ---------------------------------------------------------------------------

MODEL_FORMULATION_CHECKS = {
    "variable_definition": {
        "desc": "Variables defined with symbol, meaning, and unit",
        "patterns": [
            r'(?:其中|式中|这里)[:：]\s*\w+\s*(?:表示|为|是)',
            r'\w+\s*[:：]\s*.{2,30}(?:单位|的|表示)',
            r'符号\s*[:：]?\s*含义',
            r'变量\s*[:：]',
        ],
    },
    "model_type": {
        "desc": "Model type explicitly declared",
        "patterns": [
            r'(?:线性|非线性|多元|对数|多项式|岭|Lasso)\s*回归',
            r'(?:线性|整数|非线性|多目标|混合整数)\s*(?:规划|优化)',
            r'(?:微分方程|常微分|偏微分|ODE|PDE)',
            r'(?:分类|二分类|多分类|判定)',
            r'(?:逻辑|Logistic|SVM|支持向量机|随机森林|决策树|神经网络)',
            r'(?:聚类|K-means|层次|DBSCAN)',
            r'(?:时间序列|ARIMA|灰色预测|GM\(1[,，]1\))',
            r'(?:AHP|层次分析|TOPSIS|熵权|模糊综合评价)',
            r'(?:图论|最短路径|最小生成树|网络流)',
            r'(?:排队论|博弈论|蒙特卡洛|元胞自动机|马尔科夫)',
        ],
    },
    "assumptions": {
        "desc": "Model assumptions listed",
        "patterns": [
            r'(?:假设|前提|假定).{0,5}(?:[1-9]|[一二三四五六七八九])',
            r'(?:[（(][1-9][）)])',
        ],
    },
}

OPTIMIZATION_PATTERNS = {
    "objective": r'(?:目标函数|min\s|max\s|minimize|maximize|Z\s*=)',
    "constraints": r'(?:约束|s\.t\.|subject to|约束条件)',
}

STATISTICAL_PATTERNS = {
    "error_dist": r'(?:[ε𝜀]\s*~|正态|同方差|异方差|i\.?i\.?d|独立同分布)',
}

CLASSIFICATION_PATTERNS = {
    "decision_func": r'(?:P\s*[(（]\s*y|概率|sigmoid|判别函数|决策函数)',
    "threshold": r'(?:阈值|临界值|判定规则|判定为)',
}


def gate_model_formulation(docx_path: str, problem_num: int = 0) -> tuple[bool, str]:
    """Check model sections have proper mathematical formulation."""
    try:
        from docx import Document
    except ImportError:
        return True, "[SKIP] python-docx not available"

    path = Path(docx_path)
    if not path.exists():
        return False, f"DOCX not found: {docx_path}"

    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    issues = []
    warnings = []

    problem_markers = {"四、": "问题1", "五、": "问题2", "六、": "问题3", "七、": "问题4"}

    for marker, label in problem_markers.items():
        idx = full_text.find(marker)
        if idx == -1:
            continue
        next_idx = len(full_text)
        for m2 in list(problem_markers.keys()) + ["八、", "九、", "参考文献"]:
            mi = full_text.find(m2, idx + 2)
            if mi != -1 and mi < next_idx:
                next_idx = mi
        section_text = full_text[idx:next_idx]

        # Check 1: Variable definitions
        has_vardef = any(re.search(p, section_text) for p in MODEL_FORMULATION_CHECKS["variable_definition"]["patterns"])
        if not has_vardef:
            issues.append(f"{marker} ({label}): No variable definition found (symbol, meaning, unit)")

        # Check 2: Model type
        has_type = any(re.search(p, section_text) for p in MODEL_FORMULATION_CHECKS["model_type"]["patterns"])
        if not has_type:
            issues.append(f"{marker} ({label}): Model type not declared (e.g., '多元线性回归')")

        # Check 3: Assumptions
        assumption_count = sum(len(re.findall(p, section_text)) for p in MODEL_FORMULATION_CHECKS["assumptions"]["patterns"])
        if assumption_count < 2:
            issues.append(f"{marker} ({label}): Few model assumptions (found {assumption_count}, need >= 2)")

        # Check 4: Type-specific elements
        is_statistical = bool(re.search(r'(?:回归|R[²2]|coefficient|系数)', section_text))
        is_classification = bool(re.search(r'(?:分类|判定|Logistic|SVM|随机森林|异常)', section_text))

        if is_statistical and not re.search(STATISTICAL_PATTERNS["error_dist"], section_text):
            warnings.append(f"{marker} ({label}): Statistical model: specify error distribution assumption")

        if is_classification:
            if not re.search(CLASSIFICATION_PATTERNS["decision_func"], section_text):
                issues.append(f"{marker} ({label}): Classification model missing decision function")
            if not re.search(CLASSIFICATION_PATTERNS["threshold"], section_text):
                warnings.append(f"{marker} ({label}): Classification model: specify decision threshold")

    result_parts = []
    if issues:
        result_parts.append(f"FAILURES ({len(issues)}):")
        result_parts.extend(f"  - {i}" for i in issues)
    if warnings:
        result_parts.append(f"WARNINGS ({len(warnings)}):")
        result_parts.extend(f"  - {w}" for w in warnings)
    if not issues and not warnings:
        result_parts.append("All model formulation checks passed.")

    return len(issues) == 0, "\n".join(result_parts)



# ---------------------------------------------------------------------------
# Gate 7: Integrity Gate (Academic Integrity Verification)
# ---------------------------------------------------------------------------

INTEGRITY_CHECKS = {
    'has_references': (r'(?:????|References)', 'Must have a reference list'),
    'no_fabricated_doi': (r'doi:\\s*10\\.\\d{4,}/[^}]+', 'DOIs present (verify manually)'),
    'no_impossible_p': (r'p\\s*=\\s*0\\.000\\b', 'p=0.000 should be p<0.001'),
    'has_ci': (r'(?:????|CI|confidence)', 'Should report confidence intervals'),
    'model_type_declared': (r'??\\s*\\S*??', 'Each model must declare its type'),
    'has_symbol_table': (r'??', 'Each model should have a symbol table'),
}

def gate_integrity(docx_path: str):
    """Run academic integrity checks on assembled paper content."""
    from zipfile import ZipFile
    
    docx = Path(docx_path)
    if not docx.exists():
        return False, f'DOCX not found: {docx_path}'
    
    try:
        with ZipFile(docx, 'r') as z:
            xml_content = z.read('word/document.xml').decode('utf-8')
    except Exception as e:
        return False, f'Cannot read DOCX XML: {e}'
    
    text = re.sub(r'<[^>]+>', ' ', xml_content)
    
    issues = []
    warnings = []
    
    for check_name, (pattern, message) in INTEGRITY_CHECKS.items():
        if check_name.startswith('no_'):
            if re.search(pattern, text, re.IGNORECASE):
                warnings.append(f'WARN - {check_name}: {message}')
        else:
            if not re.search(pattern, text, re.IGNORECASE):
                if check_name in ('has_references', 'model_type_declared'):
                    issues.append(f'HARD FAIL - {check_name}: {message}')
                else:
                    warnings.append(f'WARN - {check_name}: {message}')
    
    problem_sections = len(set(re.findall(r'4\\.(\\d+)', text)))
    
    result_parts = []
    passed = len(issues) == 0
    
    if issues:
        result_parts.append(f'INTEGRITY FAILURES ({len(issues)}):')
        result_parts.extend(f'  - {i}' for i in issues)
    if warnings:
        result_parts.append(f'INTEGRITY WARNINGS ({len(warnings)}):')
        result_parts.extend(f'  - {w}' for w in warnings)
    if not issues and not warnings:
        result_parts.append('All integrity checks passed.')
    
    result_parts.insert(0, '=== ACADEMIC INTEGRITY GATE ===')
    result_parts.append(f'Sections detected: {problem_sections}')
    result_parts.append(f'Overall: {"PASS" if passed else "FAIL"}')
    
    return passed, '\n'.join(result_parts)


def main():
    p = argparse.ArgumentParser(
        description="Math Modeling Contest - Quality Gate",
    )
    sub = p.add_subparsers(dest="gate")

    pv = sub.add_parser("verify", help="Parse structured verification report")
    pv.add_argument("--stage", required=True)
    pv.add_argument("--report-file", required=True)

    ps = sub.add_parser("sanity", help="Numerical sanity check")
    ps.add_argument("--stage", required=True)
    ps.add_argument("--output-file", required=True)

    pl = sub.add_parser("lit", help="Literature citation count check")
    pl.add_argument("--stage", required=True)

    pc = sub.add_parser("consist", help="Cross-problem constant consistency")
    pc.add_argument("--problems", type=int, required=True)

    pr = sub.add_parser('render', help='Check render QA PNG artifacts')
    pr.add_argument('--render-dir', default=None, dest='render_dir')

    pp = sub.add_parser('pairing', help="Model-verify script pairing check")
    pp.add_argument("--stage", required=True)

    pd = sub.add_parser("depth", help="Check paper content depth per section")
    pd.add_argument("--docx-path", required=True, dest="docx_path")

    pf = sub.add_parser("formula", help="Check formulas present in each problem section")
    pf.add_argument("--docx-path", required=True, dest="docx_path")

    pfc = sub.add_parser("figure_ctx", help="Check figures have explanatory text")
    pfc.add_argument("--docx-path", required=True, dest="docx_path")

    pmq = sub.add_parser("model_quality", help="Check model metrics meet thresholds")
    pmq.add_argument("--report-text", required=True, dest="report_text")

    pi = sub.add_parser('integrity', help='Academic integrity verification gate')
    pi.add_argument('--docx-path', required=True, dest='docx_path')
    pmq.add_argument("--model-type", default="regression", dest="model_type")


    # KyrieZhang329-inspired gates (v5.7.0)
    pc2 = sub.add_parser("contracts", help="Print all G1-G6 gate contracts")
    ppoc = sub.add_parser("g2_poc", help="G2 PoC hard gate check")
    ppoc.add_argument("--poc-dir", default=None)
    pstale = sub.add_parser("g4_stale", help="G4 frozen staleness check")
    pstale.add_argument("--subquestion", "-q", default=None)
    pg6 = sub.add_parser("g6_audit", help="G6 enhanced 3-layer audit")
    pg6.add_argument("--workspace-path", "-w", default=None, dest="workspace_path")
    pprop = sub.add_parser("propagate", help="P1 change propagation check")
    pprop.add_argument("--changed-files", "-c", nargs="+", default=[], dest="changed_files")

    pa = sub.add_parser("all", help="Run all applicable gates")
    pa.add_argument("--stage", required=True)
    pa.add_argument("--report-file", default="")
    pa.add_argument("--output-file", default="")
    pa.add_argument("--problem-n", type=int, default=0, dest="problem_n")
    pa.add_argument("--problems", type=int, default=1)
    pa.add_argument("--docx-path", default="")
    pa.add_argument("--report-text", default="")
    pa.add_argument("--model-type", default="regression")


    args = p.parse_args()

    if args.gate in ('verify', 'sanity', 'render', 'lit', 'consist', 'pairing', 'depth', 'formula', 'figure_ctx', 'model_quality', 'model_formulation', 'integrity', 'table_formula'):
        passed = run_gate(args.gate, args)
        sys.exit(0 if passed else 1)

    elif args.gate == "all":
        results = []
        stage = args.stage

        if "_verify" in stage and args.report_file:
            results.append(run_gate("verify", args))
        if "_build" in stage and args.output_file:
            results.append(run_gate("sanity", args))
        if args.problem_n > 0:
            results.append(run_gate("lit", args))
        if args.problems > 1 and "_verify" in stage:
            results.append(run_gate("consist", args))
        if "_build" in stage:
            results.append(run_gate("pairing", args))

        if not results:
            print("[quality_gate] No applicable gates for current stage, skipping.")
            sys.exit(2)

        all_passed = all(results)
        pass_msg = chr(0x2714) + " ALL GATES PASSED" if all_passed else chr(0x2717) + " SOME GATES FAILED - cannot advance"
        print(f"\n{pass_msg}")
        sys.exit(0 if all_passed else 1)

    else:
        p.print_help()




# ============================================================
# MODEL FORMULATION GATE (v3.0) - HARD FAIL if model establishment insufficient
# ============================================================

def model_formulation_gate(doc_or_text, problem_count=3):
    """Check that paper has proper mathematical model establishment.
    
    CRITICAL: This distinguishes "model establishment" from "solution process."
    A paper that only describes algorithms without formal mathematical models FAILS.
    
    Returns: (passed, score, issues_list)
    """
    issues = []
    score = 100
    
    # Convert Document to text if needed
    if hasattr(doc_or_text, 'paragraphs'):
        text = '\n'.join(p.text for p in doc_or_text.paragraphs if p.text)
    else:
        text = str(doc_or_text)
    
    # === CHECK 1: Model type explicitly declared (15 pts) ===
    model_types = ['线性规划', '整数规划', '动态规划', '非线性规划', '多目标优化',
                   '线性回归', '非线性回归', '多元回归', '逻辑回归',
                   '微分方程', '偏微分方程', '常微分方程',
                   '分类模型', '聚类模型', '判别模型',
                   '评价模型', '层次分析', 'TOPSIS', '熵权法',
                   '预测模型', '时间序列', '灰色预测', 'ARIMA',
                   '物理模型', '光学模型', '干涉模型', 'Fabry-Perot',
                   '双光束干涉', '多光束干涉', '薄膜干涉',
                   '图论模型', '网络流', '最短路径',
                   '概率模型', '统计模型', '蒙特卡洛']
    
    found_types = [t for t in model_types if t in text]
    if len(found_types) < 1:
        issues.append("HARD FAIL: 未声明任何数学模型类型（如'建立线性回归模型'而非'建立数学模型'）")
        score -= 30
    elif len(found_types) < problem_count:
        issues.append(f"WARNING: 声明了{len(found_types)}种模型类型，但题目有{problem_count}个子问题，建议每个子问题对应一个模型类型")
        score -= 10
    
    # === CHECK 2: Symbol table present (15 pts) ===
    has_symbol_table = any(kw in text for kw in ['符号说明', '符号表', '变量定义'])
    if not has_symbol_table:
        issues.append("HARD FAIL: 缺少符号说明表。每个模型必须有独立的符号表（符号/含义/单位/类型）")
        score -= 30
    
    # === CHECK 3: Model assumptions listed (10 pts) ===
    assumption_count = text.count('假设') - text.count('假设检验')
    if assumption_count < 3:
        issues.append("FAIL: 模型假设不足（<3条）。每个模型必须列出成立的前提条件")
        score -= 20
    
    # === CHECK 4: Formal mathematical formulation (25 pts) ===
    # Check for formal math notation patterns
    has_formal_model = False
    formal_patterns = [
        '目标函数', 'min', 'max', 's.t.', '约束条件',
        '决策变量', '状态变量', '自变量', '因变量',
        '控制方程', '初始条件', '边界条件',
        '误差项', '残差', 'ε ~ N',
        '干涉条件', '光程差', '相位差',
        'Fresnel', 'Snell', 'Airy',
        '判别函数', '决策规则',
    ]
    found_formal = [p for p in formal_patterns if p.lower() in text.lower()]
    if len(found_formal) < 3:
        issues.append("HARD FAIL: 缺少正式的数学表述。模型应该有：目标函数/约束条件 或 控制方程/初始条件 或 显式的数学表达式")
        score -= 50
    else:
        has_formal_model = True
    
    # === CHECK 5: Model vs Solution distinction (15 pts) ===
    solution_keywords = ['算法步骤', '步骤一', '步骤二', '第一步', '第二步',
                         '流程图', '伪代码', '求解过程', '求解算法']
    model_keywords = ['建立', '推导', '定义', '设', '令', '记', '考虑']
    
    sol_count = sum(1 for kw in solution_keywords if kw in text)
    mod_count = sum(1 for kw in model_keywords if kw in text)
    
    if sol_count > mod_count * 2:
        issues.append("WARNING: 求解描述远多于模型建立描述。应遵循'模型在前、求解在后'的原则")
        score -= 15
    
    # === CHECK 6: Variable definitions (10 pts) ===
    # Check for undefined variables (common in weak papers)
    var_pattern = any('=' in line and not any(kw in line for kw in 
                    ['表示', '定义为', '记为', '其中', '为', '单位'])
                    for line in text.split('\n') if '=' in line and len(line) > 20)
    if var_pattern:
        issues.append("FAIL: 存在疑似未定义变量。所有含'='的公式行必须有变量解释")
        score -= 10
    
    # === CHECK 7: Content length (10 pts) ===
    text_chars = len(text)
    if text_chars < 10000:
        issues.append(f"FAIL: 内容过少 ({text_chars}字)，模型建立无法充分展开。目标: 15000-20000字")
        score -= 20
    elif text_chars > 30000:
        issues.append(f"WARNING: 内容过多 ({text_chars}字)，建议精简至20000字以内（约18-20页）")
        score -= 5
    
    passed = score >= 65 and not any('HARD FAIL' in i for i in issues)
    
    return (passed, max(0, score), issues)


# Content length standards per section (v3.0)
SECTION_LENGTH_STANDARDS = {
    'abstract':        {'min_chars': 500,  'max_chars': 1000,  'target_pages': 0.8},
    'problem_restate': {'min_chars': 800,  'max_chars': 1500,  'target_pages': 1.5},
    'problem_analysis':{'min_chars': 2000, 'max_chars': 3500,  'target_pages': 3.0},
    'assumptions':     {'min_chars': 500,  'max_chars': 1000,  'target_pages': 1.0},
    'model_building':  {'min_chars': 6000, 'max_chars': 10000, 'target_pages': 9.0},
    'results':         {'min_chars': 1000, 'max_chars': 2500,  'target_pages': 2.0},
    'evaluation':      {'min_chars': 800,  'max_chars': 2000,  'target_pages': 1.5},
    'appendix':        {'min_chars': 500,  'max_chars': 1500,  'target_pages': 1.5},
}

# Total target: 15000-22000 chars, 18-22 pages
# Model Building CORE section: 6000-10000 chars (>=35% of total)
# IRON RULE: model building section MUST exist and be the largest section



# ============================================================
# GATE: PoC (Proof of Concept) Check (v4.0)
# ============================================================

def poc_gate(methods_dir="methods", problem_count=3):
    """Check that each method candidate has a PoC (proof of concept).
    
    KyrieZhang329 Gate G2: Every candidate method must have
    a <=30 line PoC with feasibility numbers. No PoC => not validated.
    
    Returns: (passed, score, issues)
    """
    import os
    from pathlib import Path
    
    issues = []
    methods_path = Path(methods_dir)
    
    if not methods_path.exists():
        return (False, 0, ["FAIL: methods/ directory not found. Run model_build first."])
    
    subquestions_found = 0
    poc_found = 0
    
    for qdir in sorted(methods_path.iterdir()):
        if not qdir.is_dir():
            continue
        subquestions_found += 1
        
        # Check for method candidates
        candidates_file = qdir / f"{qdir.name.lower()}_method_candidates.md"
        poc_file = qdir / f"{qdir.name.lower()}_poc.py"
        
        has_candidates = candidates_file.exists()
        has_poc = poc_file.exists()
        
        if not has_candidates:
            issues.append(f"HARD FAIL: {qdir.name}: no method candidates file")
            continue
        
        if not has_poc:
            issues.append(f"FAIL: {qdir.name}: no PoC file. G2 requires <=30 line proof of concept")
            continue
        
        # Check PoC length
        try:
            poc_lines = len(poc_file.read_text(encoding="utf-8").strip().split("\n"))
            if poc_lines > 50:
                issues.append(f"WARNING: {qdir.name}: PoC is {poc_lines} lines (target <=30, max 50)")
            poc_found += 1
        except Exception:
            issues.append(f"FAIL: {qdir.name}: cannot read PoC file")
    
    if subquestions_found == 0:
        return (False, 0, ["FAIL: No subquestion method directories found"])
    
    score = int(poc_found / max(subquestions_found, 1) * 100)
    passed = poc_found >= subquestions_found and not any("HARD FAIL" in i for i in issues)
    
    return (passed, score, issues)


# ============================================================
# GATE: Assumptions Classification Check (v4.0)
# ============================================================

def assumptions_gate(assumptions_text):
    """Check that model assumptions are properly classified.
    
    KyrieZhang329 model-assumptions-builder: assumptions must be
    categorized as necessary (hard) vs simplifying (soft),
    and global vs per-subquestion.
    
    Returns: (passed, score, issues)
    """
    issues = []
    score = 100
    
    text = str(assumptions_text)
    
    if len(text) < 100:
        return (False, 0, ["HARD FAIL: Assumptions section too short or missing"])
    
    # Check: at least 3 assumptions
    assumption_markers = text.count("??") - text.count("????")
    if assumption_markers < 3:
        issues.append("FAIL: < 3 assumptions. Each model must list its preconditions")
        score -= 30
    
    # Check: assumptions have justification
    justification_keywords = ["??", "??", "???", "??", "??", "???", "???"]
    has_justification = any(kw in text for kw in justification_keywords)
    if not has_justification:
        issues.append("WARNING: Assumptions lack justification. Each should explain WHY it's reasonable")
        score -= 15
    
    # Check: distinction between necessary and simplifying (v4.0)
    has_necessary = any(kw in text for kw in ["??", "??", "??", "????", "????"])
    has_simplifying = any(kw in text for kw in ["??", "??", "????", "??", "??"])
    
    if not has_necessary:
        issues.append("WARNING: No necessary (hard) assumptions identified. Which assumptions would break the model if violated?")
        score -= 10
    if not has_simplifying:
        issues.append("INFO: No simplifying (soft) assumptions identified. Consider noting where approximations are made")
        # No score penalty - some models genuinely don't need soft assumptions
    
    # Check: global vs per-subquestion distinction
    has_global = any(kw in text for kw in ["??", "?????", "??", "??"])
    has_per_question = any(kw in text for kw in ["Q1", "Q2", "???", "???", "???"])
    
    if not has_global and not has_per_question:
        issues.append("INFO: Assumptions scope not specified. Consider marking global vs per-subquestion assumptions")
    
    passed = score >= 60 and not any("HARD FAIL" in i for i in issues)
    
    return (passed, max(0, score), issues)


# ============================================================
# GATE: Frozen Numbers Check (v4.0)
# ============================================================

def frozen_numbers_gate(frozen_dir="CUMCM_Workspace/frozen", problem_count=3):
    """Check that frozen numbers exist for all subquestions.
    
    KyrieZhang329 Gate G4: results/Qx/reports/frozen_numbers.json must exist
    and be newer than every source file.
    
    Returns: (passed, score, issues)
    """
    import json
    from pathlib import Path
    
    issues = []
    frozen_path = Path(frozen_dir)
    
    if not frozen_path.exists():
        return (False, 0, [
            "HARD FAIL: No frozen numbers directory. Run frozen_numbers.py freeze first.",
            "ACTION: python scripts/frozen_numbers.py freeze --subquestion Q1 --source results/Q1/"
        ])
    
    frozen_count = 0
    thawed_count = 0
    
    for qdir in sorted(frozen_path.iterdir()):
        if not qdir.is_dir():
            continue
        
        freeze_file = qdir / "frozen_numbers.json"
        if not freeze_file.exists():
            issues.append(f"FAIL: {qdir.name}: no frozen_numbers.json")
            continue
        
        try:
            with open(freeze_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            status = data.get("status", "frozen")
            if status == "defrosted":
                issues.append(f"FAIL: {qdir.name}: numbers are THAWED. Re-freeze after changes.")
                thawed_count += 1
            else:
                frozen_count += 1
                
        except (json.JSONDecodeError, KeyError) as e:
            issues.append(f"FAIL: {qdir.name}: corrupted frozen_numbers.json ({e})")
    
    if frozen_count == 0 and thawed_count == 0:
        return (False, 0, ["FAIL: No valid frozen numbers found for any subquestion"])
    
    score = 100 if thawed_count == 0 else max(0, 100 - thawed_count * 30)
    passed = thawed_count == 0 and frozen_count >= 1
    
    return (passed, score, issues)


# ============================================================
# GATE: G5.1 Word Count Floor (v4.0)
# ============================================================

# Word count floors per section type (Chinese chars)
SECTION_WORD_FLOORS = {
    "abstract": 200,
    "problem_restatement": 300,
    "problem_analysis": 400,
    "assumptions": 200,
    "symbols": 150,
    "data_preprocessing": 300,
    "model_construction": 600,  # per subquestion
    "model_solution": 250,      # per subquestion
    "results_analysis": 500,    # per subquestion
    "method_selection": 300,    # per subquestion
    "robustness": 350,
    "strengths_limitations": 300,
    "conclusion": 300,
}

def g5_word_count_gate(sections_dict):
    """Check that each paper section meets G5.1 word count floor.
    
    KyrieZhang329 G5.1: each section type has a minimum substantive length.
    Below floor => under_floor, not "done".
    
    sections_dict: {section_type: text_content, ...}
    
    Returns: (passed, floors_met, floors_missed)
    """
    floors_met = []
    floors_missed = []
    
    for section_type, text in sections_dict.items():
        floor = SECTION_WORD_FLOORS.get(section_type, 200)
        char_count = len(str(text))
        
        if char_count >= floor:
            floors_met.append((section_type, char_count, floor))
        else:
            floors_missed.append((section_type, char_count, floor))
    
    passed = len(floors_missed) == 0
    return (passed, floors_met, floors_missed)


# ============================================================
# GATE: G5.2 Three-Dimension Discussion Check (v4.0)
# ============================================================

DISCUSSION_DIMENSIONS = [
    "sensitivity",       # Sensitivity/robustness discussion
    "physical_meaning",  # Physical/domain meaning
    "baseline_comparison", # Baseline comparison
    "cross_consistency",  # Cross-subquestion consistency
    "uncertainty",       # Uncertainty/confidence interval
]

DIMENSION_KEYWORDS = {
    "sensitivity": ["敏感", "扰动", "鲁棒", "稳健", "perturbation", "sensitivity", "stable", "stability",
                    "change", "vary", "variation", "fluctuation"],
    "physical_meaning": ["意味着", "说明", "反映", "实际", "合理", "偏大", "偏小",
                         "means", "indicates", "reflects", "plausible", "reasonable", "requirement"],
    "baseline_comparison": ["对比", "基线", "baseline", "相比", "提升", "降低", "优于", "不如",
                            "compare", "improvement", "reduction", "better", "worse", "outperforms"],
    "cross_consistency": ["一致", "吻合", "Q1", "Q2", "Q3", "关联", "子问题",
                          "consistent", "consistency", "cross", "agreement"],
    "uncertainty": ["置信", "误差", "标准误", "区间", "CI", "不确定", "方差",
                    "confidence", "interval", "uncertainty", "standard error", "bootstrap", "range"],
}

def g5_discussion_gate(results_text):
    """Check that numerical results have >=3 discussion dimensions.
    
    KyrieZhang329 G5.2: every numerical result must be discussed from
    at least 3 of 5 dimensions.
    
    Returns: (passed, dimensions_found, issues)
    """
    text = str(results_text).lower()
    issues = []
    
    dimensions_found = {}
    for dim, keywords in DIMENSION_KEYWORDS.items():
        found_kw = [kw for kw in keywords if kw.lower() in text]
        if found_kw:
            dimensions_found[dim] = found_kw
    
    dim_count = len(dimensions_found)
    
    if dim_count < 2:
        issues.append(f"HARD FAIL: Only {dim_count}/5 discussion dimensions found. Need >=3 per numerical result.")
    elif dim_count < 3:
        issues.append(f"WARNING: Only {dim_count}/5 discussion dimensions found. Target >=3.")
    
    passed = dim_count >= 3
    
    return (passed, dimensions_found, issues)


def gate_table_formula(docx_path):
    """Check that table cells with $...$ are converted to OMML, not left as raw text.
    HARD FAIL if >5% of formula-looking table cells are unconverted."""
    import zipfile, re
    
    with zipfile.ZipFile(docx_path, "r") as z:
        xml = z.read("word/document.xml").decode("utf-8")
    
    # Find table cells
    tcs = re.findall(r"<w:tc>(.*?)</w:tc>", xml, re.DOTALL)
    
    raw_dollar = 0
    omml_cells = 0
    total_formula_cells = 0
    
    for tc in tcs:
        has_raw = bool(re.search(r"<w:t[^>]*>\$[^<]*\$</w:t>", tc))
        has_omml = "<m:oMath" in tc
        
        if has_raw:
            raw_dollar += 1
        if has_omml:
            omml_cells += 1
        if has_raw or has_omml:
            total_formula_cells += 1
    
    issues = []
    if total_formula_cells == 0:
        return True, "No formula cells in tables (OK if paper has no table formulas)"
    
    fail_rate = raw_dollar / total_formula_cells if total_formula_cells > 0 else 0
    
    if fail_rate > 0.05:
        return False, f"Table formula FAIL: {raw_dollar}/{total_formula_cells} cells have unconverted formulas ({fail_rate*100:.0f}%)"
    elif raw_dollar > 0:
        return True, f"Table formula WARN: {raw_dollar}/{total_formula_cells} unconverted (under 5% threshold, {omml_cells} OK)"
    else:
        return True, f"Table formula OK: {omml_cells} OMML cells, 0 unconverted"



# ============================================================
# G6: Three-Layer Audit (from MathModeling-skills)
# ============================================================

def audit_consistency(workspace_path: str) -> dict:
    """Layer 1: Cross-media consistency check.
    Are numbers in paper consistent with code output?
    Are symbols consistent across sections?
    """
    issues = []
    ws = Path(workspace_path)
    
    # Check if paper exists
    paper_files = list(ws.glob("output/*.docx")) + list(ws.glob("output/*.pdf"))
    if not paper_files:
        issues.append("NO_PAPER: No output paper found")
    
    # Check if CODE_MAP exists
    if not (ws / "CODE_MAP.md").exists():
        issues.append("NO_CODE_MAP: CODE_MAP.md not generated")
    
    # Check for frozen numbers
    if not list(ws.glob("output/*results*.json")):
        issues.append("NO_FROZEN: No frozen results JSON (numbers may drift)")
    
    return {"passed": len(issues) == 0, "issues": issues}


def audit_completeness(workspace_path: str) -> dict:
    """Layer 2: Completeness check.
    Does every sub-question have: model code + results + report?
    """
    issues = []
    ws = Path(workspace_path)
    
    # Check robustness reports
    rob_dir = ws / "robustness"
    if rob_dir.exists():
        for q in ["Q1", "Q2", "Q3", "Q4"]:
            rpt = rob_dir / q / f"{q.lower()}_robustness_report.md"
            if not rpt.exists():
                issues.append(f"MISSING_ROBUSTNESS: {q} has no robustness report")
    
    # Check for at least one .py solving file per contest
    py_files = list(ws.glob("output/*.py"))
    if not py_files:
        issues.append("NO_CODE: No Python solving code found")
    
    return {"passed": len(issues) <= 1, "issues": issues}


def audit_quality(workspace_path: str) -> dict:
    """Layer 3: Quality audit.
    Full workflow coherence: model depth, verification evidence, anti-fabrication.
    """
    issues = []
    ws = Path(workspace_path)
    
    # Check model depth
    codemap = ws / "CODE_MAP.md"
    if codemap.exists():
        text = codemap.read_text(encoding="utf-8")
        if len(text) < 200:
            issues.append("THIN_CODE_MAP: CODE_MAP too brief")
    
    # Check paper size (too small = likely incomplete)
    for docx in ws.glob("output/*.docx"):
        if docx.stat().st_size < 30000:  # < 30KB
            issues.append(f"SMALL_PAPER: {docx.name} may be incomplete")
    
    return {"passed": len(issues) == 0, "issues": issues}


def run_g6_audit(workspace_path: str) -> dict:
    """Run full G6 three-layer audit"""
    results = {
        "consistency": audit_consistency(workspace_path),
        "completeness": audit_completeness(workspace_path),
        "quality": audit_quality(workspace_path),
    }
    all_passed = all(r["passed"] for r in results.values())
    results["gate_passed"] = all_passed
    return results


def print_g6_report(results: dict):
    """Print G6 audit report"""
    print("\n" + "="*60)
    print("  G6 Audit Layer (Consistency -> Completeness -> Quality)")
    print("="*60)
    
    icons = {True: "PASS", False: "FAIL"}
    for layer, r in results.items():
        if layer == "gate_passed":
            continue
        icon = icons[r["passed"]]
        print(f"\n  [{icon}] {layer.upper()}")
        for issue in r.get("issues", []):
            print(f"    - {issue}")
    
    gate_icon = "PASS" if results["gate_passed"] else "BLOCKED"
    print(f"\n  G6 GATE: [{gate_icon}]")
    if not results["gate_passed"]:
        print("  Pipeline blocked until all issues resolved.")
    print("="*60)

from pathlib import Path

if __name__ == "__main__":
    main()
