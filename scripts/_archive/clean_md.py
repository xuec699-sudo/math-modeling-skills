# -*- coding: utf-8 -*-
"""Clean markdown artifacts from DOCX"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")

cleaned_count = 0
removed_count = 0

for para in doc.paragraphs:
    text = para.text.strip()
    
    # Case 1: "#### Text" headings → strip "#### " prefix
    if text.startswith('#### '):
        new_text = text[5:].strip()  # Remove "#### "
        # Clear and rewrite
        for run in para.runs:
            run._element.getparent().remove(run._element)
        # Add cleaned text
        run = para.add_run(new_text)
        # Try to make it bold to look like a heading
        run.bold = True
        run.font.size = Pt(12)
        cleaned_count += 1
        print(f"  [FIXED] '#### ' removed: '{new_text[:60]}'")
        continue
    
    # Case 2: Pure "---" separator → remove the paragraph
    if text == '---':
        # Remove from document
        para._element.getparent().remove(para._element)
        removed_count += 1
        continue
    
    # Case 3: Pure "```" code fences → remove
    if text == '```':
        para._element.getparent().remove(para._element)
        removed_count += 1
        continue

# Save
out = r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx"
doc.save(out)

print(f"\nCleaned: {cleaned_count} headings fixed, {removed_count} separators removed")
print(f"Total: {cleaned_count + removed_count} artifacts cleaned")

# Quick re-verify
doc2 = Document(out)
remaining = 0
for para in doc2.paragraphs:
    t = para.text.strip()
    if t.startswith('#### ') or t == '---' or t == '```':
        remaining += 1
        print(f"  [WARN] Still found: '{t[:60]}'")

if remaining == 0:
    print(f"VERIFIED: 0 markdown artifacts remaining")
else:
    print(f"WARNING: {remaining} artifacts still present")

import os
print(f"File size: {os.path.getsize(out)/1024:.0f} KB")
print(f"\n[DONE] DOCX cleaned")
