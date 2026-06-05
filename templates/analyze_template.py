import sys; sys.stdout.reconfigure(encoding="utf-8")
from docx import Document, shared
from docx.oxml.ns import qn

doc = Document("2026建模竞赛A题_最终论文.docx")

# Page setup
print("=== PAGE SETUP ===")
for s in doc.sections:
    print(f"Page: {s.page_width.cm:.1f} x {s.page_height.cm:.1f} cm")
    print(f"Margins: T={s.top_margin.cm:.1f} B={s.bottom_margin.cm:.1f} L={s.left_margin.cm:.1f} R={s.right_margin.cm:.1f} cm")

# Styles
print("\n=== STYLES ===")
styles = {}
for p in doc.paragraphs:
    sname = p.style.name if p.style else "None"
    if sname not in styles:
        styles[sname] = {"count": 0, "sample": "", "sz": None, "bold": None, "align": None}
    styles[sname]["count"] += 1
    if not styles[sname]["sample"] and p.text.strip():
        styles[sname]["sample"] = p.text.strip()[:60]
    for r in p.runs:
        if r.font.size:
            styles[sname]["sz"] = r.font.size.pt
        if r.font.bold:
            styles[sname]["bold"] = True
        if p.alignment:
            styles[sname]["align"] = str(p.alignment)

for sn, info in sorted(styles.items(), key=lambda x: -x[1]["count"]):
    print(f"  {sn}: {info['count']}x, sz={info['sz']}, bold={info['bold']}, align={info['align']}  |{info['sample']}")

# First paragraphs
print("\n=== FIRST 20 PARA ===")
for i, p in enumerate(doc.paragraphs[:20]):
    t = p.text.strip()[:80]
    sn = p.style.name
    sz = None
    for r in p.runs:
        if r.font.size:
            sz = r.font.size.pt; break
    print(f"  P{i}: [{sn}] sz={sz} | {t}")

# Table analysis
print(f"\n=== TABLES: {len(doc.tables)} ===")
for ti, tbl in enumerate(doc.tables[:3]):
    print(f"\nTable {ti+1}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
    # Header
    for j, cell in enumerate(tbl.rows[0].cells):
        for r in cell.paragraphs[0].runs:
            print(f"  Header[{j}]: sz={r.font.size.pt if r.font.size else '?'} bold={r.font.bold} font={r.font.name} | {cell.text[:30]}")
    # First data row
    if len(tbl.rows) > 1:
        for j, cell in enumerate(tbl.rows[1].cells):
            for r in cell.paragraphs[0].runs:
                print(f"  Data[{j}]: sz={r.font.size.pt if r.font.size else '?'} bold={r.font.bold} font={r.font.name} | {cell.text[:30]}")
    # Check borders
    tbl_pr = tbl._tbl.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr")
    if tbl_pr is not None:
        borders = tbl_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders")
        if borders is not None:
            print("  Borders:", {b.tag.split("}")[1]: f"val={b.get(qn('w:val'))} sz={b.get(qn('w:sz'))}" for b in borders})
