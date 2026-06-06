# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")
all_text = []
for para in doc.paragraphs:
    t = para.text.strip()
    if t:
        all_text.append(t)

# Also extract tables
for i, table in enumerate(doc.tables):
    all_text.append(f'\n**表{i+1}**')
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        all_text.append(' | '.join(cells))

full = '\n\n'.join(all_text)
out = r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_text_r3.txt"
with open(out, 'w', encoding='utf-8') as f:
    f.write(full)
print(f"Extracted: {len(full):,} chars, saved to paper_text_r3.txt")
