# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt

doc = Document(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")

cleaned = 0
removed = 0
to_remove = []

for para in doc.paragraphs:
    text = para.text.strip()
    
    if text.startswith('#### '):
        new_text = text[5:].strip()
        for run in para.runs:
            run._element.getparent().remove(run._element)
        run = para.add_run(new_text)
        run.bold = True
        run.font.size = Pt(12)
        cleaned += 1
        continue
    
    if text == '---' or text == '```':
        to_remove.append(para)
        removed += 1
        continue

# Remove separator paragraphs
for para in to_remove:
    para._element.getparent().remove(para._element)

# Save to new file
out = r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final_clean.docx"
doc.save(out)

# Verify
doc2 = Document(out)
artifacts = 0
for para in doc2.paragraphs:
    t = para.text.strip()
    if t.startswith('#### ') or t == '---' or t == '```':
        artifacts += 1

import zipfile, re, os
with zipfile.ZipFile(out, "r") as z:
    xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    images = [n for n in z.namelist() if "media" in n and n.endswith('.png')]
    omml = len(re.findall(r"<m:oMath", xml))

print(f"Cleaned: {cleaned} #### headings, {removed} separators removed")
print(f"Remaining artifacts: {artifacts}")
print(f"Images: {len(images)}, OMML: {omml}")
print(f"Size: {os.path.getsize(out)/1024:.0f} KB")
print(f"\nSaved: paper_final_clean.docx")
print(f"Close paper_final.docx in Word, then rename paper_final_clean.docx -> paper_final.docx")
