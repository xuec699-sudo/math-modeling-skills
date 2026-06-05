# -*- coding: utf-8 -*-
"""Extract text from updated DOCX for review"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import re

doc = Document(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")

# Extract all paragraph text
all_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        # Detect heading style
        if para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '')
            prefix = '#' * int(level) + ' '
            all_text.append(prefix + text)
        else:
            all_text.append(text)

full_text = '\n\n'.join(all_text)

# Extract tables too
for i, table in enumerate(doc.tables):
    all_text.append(f'\n**表{i+1}**')
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        all_text.append(' | '.join(cells))

full_text = '\n\n'.join(all_text)

# Stats
chars = len(full_text)
omml_count = sum(1 for p in doc.paragraphs if '<m:oMath' in p._element.xml)

print(f"Extracted text: {chars:,} chars")
print(f"OMML equations in DOCX: {omml_count}")

# Save for review
out = r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_text_for_review.txt"
with open(out, 'w', encoding='utf-8') as f:
    f.write(full_text)

print(f"Saved to: {out}")
