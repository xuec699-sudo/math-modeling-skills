# -*- coding: utf-8 -*-
"""Fix reference order in DOCX"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

doc = Document(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")

# Find paragraphs with [10], [11], [12] references
ref_paras = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith('[10]') or text.startswith('[11]') or text.startswith('[12]'):
        ref_paras.append((i, text[:60]))

print("Found reference paragraphs:")
for idx, txt in ref_paras:
    print(f"  [{idx}] {txt}")

# If they're in wrong order, swap the text content
if len(ref_paras) == 3:
    # Get the current order
    texts = [(i, doc.paragraphs[i].text) for i, _ in ref_paras]
    
    # Sort by reference number
    import re
    def get_ref_num(text):
        m = re.match(r'\[(\d+)\]', text.strip())
        return int(m.group(1)) if m else 99
    
    sorted_texts = sorted(texts, key=lambda x: get_ref_num(x[1]))
    
    print(f"\nCurrent order: {[t[0] for t in texts]}")
    print(f"Desired order: {[t[0] for t in sorted_texts]}")
    
    if [t[0] for t in texts] != [t[0] for t in sorted_texts]:
        # Swap the text between paragraphs to fix order
        indices = [t[0] for t in texts]
        desired_indices = [t[0] for t in sorted_texts]
        
        # Store all 3 texts
        stored_texts = [doc.paragraphs[idx].text for idx in indices]
        stored_runs = []
        for idx in indices:
            runs_data = [(r.text, r.bold, r.italic) for r in doc.paragraphs[idx].runs]
            stored_runs.append(runs_data)
        
        # Clear and rewrite in correct order
        # Actually it's easier: just reassign the text
        for idx, correct_text in zip(indices, [t[1] for t in sorted_texts]):
            para = doc.paragraphs[idx]
            # Clear runs
            for run in para.runs:
                run._element.getparent().remove(run._element)
            para.add_run(correct_text)
        
        print("References reordered correctly: [10], [11], [12]")
    else:
        print("References already in correct order")
else:
    print(f"Expected 3 ref paras, found {len(ref_paras)}")

doc.save(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")
print("\n[DONE] Reference order fixed")
