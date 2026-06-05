# -*- coding: utf-8 -*-
"""Map DOCX paragraph structure for precise insertion"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")

# Find key paragraphs
targets = [
    "模型评估与特征重要性",  # 4.1.4 heading
    "两个模型在训练集上",     # model eval paragraph
    "随机森林模型输出的特征重要性", # feature importance
    "模型超参数",            # check if already exists
    "模型二",               # 4.2 heading
    "喷头网格化布设",        # 4.2.2
    "总计21个喷头",         # near end of 4.2.2
    "参考文献",             # references
    "孙景生",              # last reference
    "附录A",               # appendix
    "供水-需求平衡方程",    # 4.3.2
]

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    for target in targets:
        if target in text:
            style = para.style.name if para.style else "None"
            # Show context: prev and next paragraph snippets
            prev_text = doc.paragraphs[i-1].text[:50] if i > 0 else "(START)"
            next_text = doc.paragraphs[i+1].text[:50] if i+1 < len(doc.paragraphs) else "(END)"
            print(f"[{i:03d}] style={style:20s} | \"{text[:80]}\"")
            print(f"       prev: \"{prev_text}\"")
            print(f"       next: \"{next_text}\"")
            break

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
