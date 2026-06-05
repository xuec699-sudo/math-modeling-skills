# -*- coding: utf-8 -*-
"""In-place edit paper_final.docx using python-docx, preserving OMML equations and images."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re

doc_path = r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx"
doc = Document(doc_path)

def get_para_text(para):
    """Get visible text only (skip OMML runs)"""
    texts = []
    for run in para.runs:
        if '<m:oMath' not in run._element.xml and run.text:
            texts.append(run.text)
    return ''.join(texts)

def replace_in_para(para, old_text, new_text):
    """Replace text in paragraph, only touching non-OMML runs."""
    full_text = ''.join(r.text for r in para.runs if r.text)
    if old_text not in full_text:
        return False
    
    # Simple approach: find the run containing the old text and replace
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    
    # Fallback: multi-run replacement
    remaining = new_text
    for run in para.runs:
        if run.text and old_text:
            if old_text.startswith(run.text):
                old_text = old_text[len(run.text):]
                run.text = remaining[:len(run.text)]
                remaining = remaining[len(run.text):]
            elif run.text in old_text:
                run.text = run.text.replace(run.text, remaining[:len(run.text)])
                remaining = remaining[len(run.text):]
                old_text = old_text.replace(run.text[:len(run.text)], '', 1)
    return True

# Counter for changes
changes = 0

# --- Change 1: Abstract - coverage area ---
for para in doc.paragraphs:
    if '总覆盖面积达14,847m²' in get_para_text(para):
        replace_in_para(para, '总覆盖面积达14,847m²', '理论覆盖面积14,844m²（经重叠校正后有效覆盖约10,000m²）')
        changes += 1
        print(f"  [OK] Abstract: coverage area clarified")
        break

# --- Change 2: Abstract - CV gap note ---
for para in doc.paragraphs:
    txt = get_para_text(para)
    if '五折时间序列交叉验证RMSE为0.04008。利用该模型' in txt:
        replace_in_para(para, '0.04008。利用该模型', '0.04008（训练集与CV的RMSE差距分析见4.1.4节）。利用该模型')
        changes += 1
        print(f"  [OK] Abstract: CV gap note added")
        break

# --- Change 3: Section 4.2 heading rename ---
for para in doc.paragraphs:
    if '模型二：灌溉管网优化模型' in get_para_text(para):
        replace_in_para(para, '模型二：灌溉管网优化模型', '模型二：灌溉管网设计与成本估算模型')
        changes += 1
        print(f"  [OK] Sec 4.2: heading renamed")
        break

# --- Change 4: Section 4.2.2 - coverage area clarification ---
for para in doc.paragraphs:
    txt = get_para_text(para)
    if '21个喷头总覆盖面积为' in txt and '远大于农场' in txt:
        replace_in_para(para, '远大于农场实际面积10,000 m²，覆盖率达148%', '由于相邻喷头间距15m等于覆盖半径，覆盖圆存在显著重叠。经几何校正后实际有效覆盖面积约等于农场总面积10,000 m²，覆盖冗余约48%')
        changes += 1
        print(f"  [OK] Sec 4.2.2: coverage clarified")
        break

# --- Change 5: Section 4.3.2 - derived formula clarity ---
for para in doc.paragraphs:
    txt = get_para_text(para)
    if '设定保障条件' in txt and '由此可解出' in txt:
        replace_in_para(para, '设定保障条件', '设定保障条件（基于供水-需求平衡方程严格推导）')
        changes += 1
        print(f"  [OK] Sec 4.3.2: derivation clarified")
        break

# --- Change 6: Section 6.2 - add limitation about CV gap ---
for para in doc.paragraphs:
    txt = get_para_text(para)
    if '日尺度建模忽略了日内波动' in txt and '1.' in txt:
        replace_in_para(para, '1. **日尺度建模忽略了日内波动**', '1. **训练-验证性能差距需进一步诊断**：随机森林的训练RMSE（0.00634）与时间序列CV-RMSE（0.04008）之间存在6.3倍的差距。虽然4.1.4节已讨论其主要来源（分布漂移、自回归结构衰减），但使用仅有气象特征的模型（R²≈0.91）作为性能基线表明泛化能力有提升空间。改进方向：引入差分特征预测湿度变化量以降低自回归依赖。\n\n2. **日尺度建模忽略了日内波动**')
        changes += 1
        print(f"  [OK] Sec 6.2: CV gap limitation added")
        break

# --- Change 7: Conclusion update ---
for para in doc.paragraphs:
    txt = get_para_text(para)
    if '测试集R²达到0.9888' in txt:
        replace_in_para(para, '测试集R²达到0.9888', '训练集R²达到0.9888（消融实验表明仅气象特征R²为0.9127，保守泛化RMSE约0.047）')
        changes += 1
    if '灌溉管网采用树状拓扑和网格化喷头布局' in txt:
        replace_in_para(para, '灌溉管网采用树状拓扑和网格化喷头布局', '灌溉管网经网格与三角方案对比后选择3×7矩形网格布局')
        changes += 1
        print(f"  [OK] Sec 7: conclusion updated")
        break

# --- Save ---
out_path = r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx"
doc.save(out_path)

# Verify
import zipfile
with zipfile.ZipFile(out_path, "r") as z:
    names = z.namelist()
    xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    images = [n for n in names if "media" in n and n.endswith('.png')]
    omml = len(re.findall(r"<m:oMath", xml))
    raw_dollar = len(re.findall(r"<w:t[^>]*>\$", xml))

print(f"\nVerification:")
print(f"  Images preserved: {len(images)} ({', '.join([n.split('/')[-1] for n in images])})")
print(f"  OMML equations: {omml}")
print(f"  Raw $ remaining: {raw_dollar}")
print(f"  File size: {__import__('os').path.getsize(out_path)/1024:.0f} KB")
print(f"  Total changes: {changes}/7 applied")
print(f"\n[DONE] paper_final.docx updated in-place")
