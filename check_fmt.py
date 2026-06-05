import sys; sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
doc = Document("CUMCM_Workspace/output/paper_template_format.docx")

s = doc.sections[0]
print(f"Page: {s.page_width.cm:.1f}x{s.page_height.cm:.1f}cm")
print(f"Margins: T={s.top_margin.cm:.1f} B={s.bottom_margin.cm:.1f} L={s.left_margin.cm:.1f} R={s.right_margin.cm:.1f}cm")

if doc.tables:
    tbl = doc.tables[0]
    for i in range(min(2, len(tbl.rows))):
        for j in range(min(2, len(tbl.rows[i].cells))):
            cell = tbl.rows[i].cells[j]
            for r in cell.paragraphs[0].runs:
                sz = r.font.size.pt if r.font.size else "inherit"
                print(f"Tbl r{i}c{j}: sz={sz}, bold={r.font.bold}, font={r.font.name} | {r.text[:20]}")
                break

# Check heading
for p in doc.paragraphs:
    if p.style.name.startswith("Heading") and p.text.strip():
        for r in p.runs:
            print(f"{p.style.name}: sz={r.font.size.pt if r.font.size else '?'}, bold={r.font.bold} | {p.text[:40]}")
            break
        break
