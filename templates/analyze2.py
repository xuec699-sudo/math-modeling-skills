import sys; sys.stdout.reconfigure(encoding="utf-8")
from docx import Document

doc = Document("2026建模竞赛A题_最终论文.docx")

# Check "First Paragraph" style details
fp_style = doc.styles["First Paragraph"]
print("=== First Paragraph style ===")
print(f"  Font: {fp_style.font.name}, size={fp_style.font.size}")
print(f"  Line spacing: {fp_style.paragraph_format.line_spacing}")
print(f"  First line indent: {fp_style.paragraph_format.first_line_indent}")
print(f"  Space before: {fp_style.paragraph_format.space_before}")
print(f"  Space after: {fp_style.paragraph_format.space_after}")

# Check Normal style
n_style = doc.styles["Normal"]
print("\n=== Normal style ===")
print(f"  Font: {n_style.font.name}, size={n_style.font.size}")

# Check Heading styles
for lvl in [1, 2, 3]:
    h = doc.styles[f"Heading {lvl}"]
    print(f"\n=== Heading {lvl} ===")
    print(f"  Font: {h.font.name}, size={h.font.size}, bold={h.font.bold}")
    rPr = h.element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if rPr is not None:
        rF = rPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
        if rF is not None:
            print(f"  EastAsia font: {rF.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')}")

# Check figure caption example
print("\n=== Figure caption sample ===")
for p in doc.paragraphs:
    if p.text.strip().startswith("图") and "流程图" in p.text:
        print(f"  Style: {p.style.name}")
        print(f"  Text: {p.text[:60]}")
        for r in p.runs:
            print(f"  Run: sz={r.font.size}, bold={r.font.bold}, font={r.font.name}")
        break

# Check the table caption style
tc_style = doc.styles["Table Caption"]
print(f"\n=== Table Caption style ===")
print(f"  Font: {tc_style.font.name}, size={tc_style.font.size}, bold={tc_style.font.bold}")

# Body text sample
print("\n=== Body text sample (First Paragraph) ===")
for p in doc.paragraphs:
    if p.style.name == "First Paragraph" and len(p.text) > 100:
        print(f"  Text: {p.text[:80]}...")
        for r in p.runs[:2]:
            print(f"  Run: sz={r.font.size}, font={r.font.name}")
        break
