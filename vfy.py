import sys; sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
doc = Document("CUMCM_Workspace/output/paper_final.docx")
s = doc.sections[0]
print(f"Page: {s.page_width.cm:.1f}x{s.page_height.cm:.1f}cm, M: T{s.top_margin.cm:.1f}B{s.bottom_margin.cm:.1f}L{s.left_margin.cm:.1f}R{s.right_margin.cm:.1f}")
if doc.tables:
    t = doc.tables[1]
    for ri in range(min(2, len(t.rows))):
        for ci in range(min(2, len(t.rows[ri].cells))):
            runs = t.rows[ri].cells[ci].paragraphs[0].runs
            for rn, r in enumerate(runs):
                if r.text.strip():
                    sz = r.font.size.pt if r.font.size else "?"
                    print(f"T[{ri},{ci}] run{rn}: sz={sz}, bold={r.font.bold}, font={r.font.name} | {r.text[:25]}")
print("Format check done")
