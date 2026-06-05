# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r"D:\Codex\skills\math-modeling-contest\CUMCM_Workspace\output\paper_final.docx")

# Find paragraphs with markdown syntax artifacts
md_patterns = ['---', '####', '###', '## ', '```', '**', '__']

print("Paragraphs with Markdown artifacts:")
print("=" * 60)
found = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    for pat in md_patterns:
        if pat in text and text != pat:  # skip pure separators
            print(f"[{i}] pattern='{pat}' | \"{text[:100]}\"")
            found += 1
            break

# Also find pure --- and #### lines
print(f"\nPure markdown lines:")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text in ['---', '####', '###', '##', '#', '```']:
        print(f"[{i}] \"{text}\" (prev: \"{doc.paragraphs[i-1].text.strip()[:60] if i>0 else 'START'}\")")
        found += 1

print(f"\nTotal artifacts found: {found}")
